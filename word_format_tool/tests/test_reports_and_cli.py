from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from typer.testing import CliRunner

from word_format_tool.cli import app
from word_format_tool.report_writer import write_json_report, write_markdown_report


runner = CliRunner()


def test_cli_exposes_structure_and_visual_commands() -> None:
    result = runner.invoke(app, ["--help"])

    assert result.exit_code == 0
    assert "plan-structure" in result.stdout
    assert "complete-structure" in result.stdout
    assert "visual-compare" in result.stdout
    assert "inspect-structure" in result.stdout
    assert "apply-operations" in result.stdout
    assert "validate-layout" in result.stdout
    assert "refresh-fields" in result.stdout


def test_report_writers_create_machine_and_human_readable_files(
    draft_path: Path, rules_path: Path, tmp_path: Path,
) -> None:
    from word_format_tool import inspect_document

    report = inspect_document(draft_path, rules_path)
    json_path = tmp_path / "report.json"
    md_path = tmp_path / "report.md"

    write_json_report(report, json_path)
    write_markdown_report(report, md_path)

    assert json.loads(json_path.read_text(encoding="utf-8"))["issues"]
    markdown = md_path.read_text(encoding="utf-8")
    assert "# Word 格式检查报告" in markdown
    assert "## 覆盖范围" in markdown
    assert "## 问题列表" in markdown


def test_validate_rules_cli(rules_path: Path) -> None:
    result = runner.invoke(app, ["validate-rules", str(rules_path)])

    assert result.exit_code == 0
    assert "规则校验通过" in result.stdout


def test_inspect_and_fix_cli_create_outputs(
    draft_path: Path, rules_path: Path, tmp_path: Path,
) -> None:
    inspect_json = tmp_path / "inspect.json"
    inspect_md = tmp_path / "inspect.md"
    inspect_result = runner.invoke(
        app,
        [
            "inspect",
            str(draft_path),
            "--rules",
            str(rules_path),
            "-o",
            str(inspect_json),
            "--md",
            str(inspect_md),
        ],
    )

    fixed = tmp_path / "fixed.docx"
    fix_json = tmp_path / "fix.json"
    fix_md = tmp_path / "fix.md"
    fix_result = runner.invoke(
        app,
        [
            "fix",
            str(draft_path),
            "--rules",
            str(rules_path),
            "-o",
            str(fixed),
            "--report",
            str(fix_json),
            "--md",
            str(fix_md),
        ],
    )

    assert inspect_result.exit_code == 0
    assert inspect_json.exists() and inspect_md.exists()
    assert fix_result.exit_code == 0
    assert fixed.exists() and fix_json.exists() and fix_md.exists()


def test_inspect_cli_does_not_require_markdown_output(
    draft_path: Path, rules_path: Path, tmp_path: Path,
) -> None:
    inspect_json = tmp_path / "inspect.json"

    result = runner.invoke(
        app,
        [
            "inspect",
            str(draft_path),
            "--rules",
            str(rules_path),
            "-o",
            str(inspect_json),
        ],
    )

    assert result.exit_code == 0
    assert inspect_json.exists()


def test_link_references_cli_only_changes_reference_navigation(
    tmp_path: Path,
    rules_data: dict,
) -> None:
    rules_data["page"] = None
    rules_data["roles"] = {
        "body": {"line_spacing": 1.25},
        "heading_1": {"alignment": "center"},
        "reference": {
            "line_spacing": 1.25,
            "hanging_indent_chars": 1.5,
        },
    }
    rules_path = tmp_path / "reference-rules.json"
    rules_path.write_text(
        json.dumps(rules_data, ensure_ascii=False), encoding="utf-8"
    )

    source = tmp_path / "reference-cli.docx"
    document = Document()
    document.add_paragraph("已有研究[1]。")
    document.add_paragraph("参考文献")
    entry = document.add_paragraph("[1] 第一条文献")
    entry.paragraph_format.line_spacing = 2.0
    document.save(source)

    output = tmp_path / "reference-linked.docx"
    report = tmp_path / "reference-linked.json"
    markdown = tmp_path / "reference-linked.md"
    result = runner.invoke(
        app,
        [
            "link-references",
            str(source),
            "--rules",
            str(rules_path),
            "-o",
            str(output),
            "--report",
            str(report),
            "--md",
            str(markdown),
        ],
    )

    assert result.exit_code == 0
    assert output.exists() and report.exists() and markdown.exists()
    linked = Document(output)
    assert linked.paragraphs[2].paragraph_format.line_spacing == 2.0
    assert linked.paragraphs[0]._p.xpath(
        ".//*[local-name()='hyperlink' "
        "and @*[local-name()='anchor']='_SDAU_REF_0001']"
    )

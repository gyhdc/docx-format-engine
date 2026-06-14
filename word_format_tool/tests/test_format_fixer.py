from __future__ import annotations

import base64
import json
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement

from word_format_tool import fix_document, inspect_document
from word_format_tool import document_io, format_fixer
from word_format_tool.exceptions import DocumentReadError


def test_fix_reduces_fixable_issues_and_preserves_text(
    draft_path: Path, rules_path: Path, tmp_path: Path,
) -> None:
    output = tmp_path / "fixed.docx"
    original_text = [paragraph.text for paragraph in Document(draft_path).paragraphs]
    before = inspect_document(draft_path, rules_path)

    result = fix_document(draft_path, rules_path, output)

    after = inspect_document(output, rules_path)
    assert output.exists()
    assert [paragraph.text for paragraph in Document(output).paragraphs] == original_text
    assert after["summary"]["total_issues"] < before["summary"]["total_issues"]
    assert result["summary"]["fixed_issues"] > 0


def test_fix_does_not_overwrite_original_by_default(
    draft_path: Path, rules_path: Path, tmp_path: Path,
) -> None:
    original_bytes = draft_path.read_bytes()

    fix_document(draft_path, rules_path, tmp_path / "fixed.docx")

    assert draft_path.read_bytes() == original_bytes


def test_unknown_role_is_not_force_formatted(
    draft_path: Path, rules_path: Path, tmp_path: Path,
) -> None:
    output = tmp_path / "fixed.docx"
    before = Document(draft_path).paragraphs[3]._p.xml

    fix_document(draft_path, rules_path, output)

    assert Document(output).paragraphs[3]._p.xml == before


def test_fix_rejects_same_input_and_output_path(
    draft_path: Path, rules_path: Path,
) -> None:
    import pytest

    from word_format_tool.exceptions import UnsafeOutputPathError

    with pytest.raises(UnsafeOutputPathError):
        fix_document(draft_path, rules_path, draft_path)


def test_fix_preserves_images_tables_and_formulas(
    tmp_path: Path, rules_path: Path,
) -> None:
    source = tmp_path / "objects.docx"
    image = tmp_path / "pixel.png"
    image.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwC"
            "AAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
    )
    document = Document()
    picture_paragraph = document.add_paragraph()
    picture_paragraph.add_run().add_picture(str(image))
    document.add_table(rows=1, cols=1).cell(0, 0).text = "保留表格"
    formula_paragraph = document.add_paragraph()
    formula = OxmlElement("m:oMath")
    formula_run = OxmlElement("m:r")
    formula_text = OxmlElement("m:t")
    formula_text.text = "x+y"
    formula_run.append(formula_text)
    formula.append(formula_run)
    formula_paragraph._p.append(formula)
    document.save(source)

    output = tmp_path / "objects-fixed.docx"
    fix_document(source, rules_path, output)

    fixed = Document(output)
    assert len(fixed.inline_shapes) == 1
    assert len(fixed.tables) == 1
    assert (
        fixed.element.xpath("count(.//*[local-name()='oMath'])")
        == 1.0
    )


def test_fix_only_changes_fields_reported_as_wrong(
    tmp_path: Path, rules_data: dict,
) -> None:
    rules_data["page"] = None
    rules_data["roles"] = {
        "body": {
            "bold": False,
            "line_spacing": 1.5,
        }
    }
    rules_path = tmp_path / "selective-rules.json"
    rules_path.write_text(
        json.dumps(rules_data, ensure_ascii=False), encoding="utf-8"
    )

    source = tmp_path / "selective.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("这是一段足够长的普通正文，用作代表格式。")
    emphasis = paragraph.add_run("重点")
    emphasis.bold = True
    paragraph.paragraph_format.line_spacing = 1.0
    document.save(source)

    output = tmp_path / "selective-fixed.docx"
    result = fix_document(source, rules_path, output)

    fixed = Document(output).paragraphs[0]
    assert fixed.paragraph_format.line_spacing == pytest.approx(1.5)
    assert fixed.runs[1].bold is True
    assert {
        issue["field"] for issue in result["issues"] if issue["fixed"]
    } == {"line_spacing"}


def test_fix_validation_failure_does_not_replace_existing_output(
    draft_path: Path,
    rules_path: Path,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    output = tmp_path / "existing.docx"
    original_output = b"existing output must survive"
    output.write_bytes(original_output)

    def reject_preservation(*args: object, **kwargs: object) -> None:
        raise DocumentReadError("simulated preservation failure")

    monkeypatch.setattr(
        format_fixer,
        "validate_preserved_content",
        reject_preservation,
        raising=False,
    )

    with pytest.raises(DocumentReadError, match="preservation failure"):
        fix_document(draft_path, rules_path, output)

    assert output.read_bytes() == original_output


def test_content_validation_accepts_formatting_only_and_rejects_content_changes(
    tmp_path: Path,
) -> None:
    validator = getattr(document_io, "validate_preserved_content", None)
    assert callable(validator)

    source = tmp_path / "preservation-source.docx"
    document = Document()
    paragraph = document.add_paragraph("正文内容")
    paragraph.add_run("重点").bold = True
    document.add_table(rows=1, cols=1).cell(0, 0).text = "表格内容"
    formula_paragraph = document.add_paragraph()
    formula = OxmlElement("m:oMath")
    formula_text = OxmlElement("m:t")
    formula_text.text = "x+y"
    formula.append(formula_text)
    formula_paragraph._p.append(formula)
    document.save(source)

    formatting_only = tmp_path / "formatting-only.docx"
    formatted = Document(source)
    formatted.paragraphs[0].paragraph_format.line_spacing = 1.5
    formatted.save(formatting_only)
    validator(source, formatting_only)

    changed = tmp_path / "content-changed.docx"
    changed_document = Document(source)
    changed_document.paragraphs[0].text = "正文被修改"
    changed_document.save(changed)

    with pytest.raises(DocumentReadError, match="内容完整性"):
        validator(source, changed)

    image = tmp_path / "pixel.png"
    image_bytes = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwC"
        "AAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
    )
    image.write_bytes(image_bytes)
    media_source = tmp_path / "media-source.docx"
    media_document = Document()
    media_document.add_paragraph().add_run().add_picture(str(image))
    media_document.save(media_source)

    changed_image = tmp_path / "pixel-changed.png"
    changed_image.write_bytes(image_bytes + b"safe trailing bytes")
    media_changed = tmp_path / "media-changed.docx"
    changed_media_document = Document()
    changed_media_document.add_paragraph().add_run().add_picture(
        str(changed_image)
    )
    changed_media_document.save(media_changed)

    with pytest.raises(DocumentReadError, match="media"):
        validator(media_source, media_changed)


def test_fix_without_fixable_issues_copies_source_bytes_exactly(
    tmp_path: Path,
    rules_data: dict,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    rules_data["page"] = None
    rules_data["roles"] = {"body": {}}
    rules_path = tmp_path / "no-op-rules.json"
    rules_path.write_text(
        json.dumps(rules_data, ensure_ascii=False), encoding="utf-8"
    )

    source = tmp_path / "already-correct.docx"
    document = Document()
    document.add_paragraph("无需修改的正文")
    document.save(source)

    def reject_document_reserialization(*args: object, **kwargs: object) -> None:
        raise AssertionError("no-op fixes must not reserialize the DOCX")

    monkeypatch.setattr(
        format_fixer,
        "save_docx",
        reject_document_reserialization,
        raising=False,
    )

    output = tmp_path / "already-correct-fixed.docx"
    result = fix_document(source, rules_path, output)

    assert result["summary"]["fixable_issues"] == 0
    assert output.read_bytes() == source.read_bytes()


def test_fix_formats_and_links_references_without_changing_visible_text(
    tmp_path: Path,
    rules_data: dict,
) -> None:
    rules_data["page"] = None
    rules_data["roles"] = {
        "body": {},
        "heading_1": {
            "font_size_pt": 14,
            "bold": True,
            "alignment": "center",
        },
        "reference": {
            "font_size_pt": 12,
            "alignment": "both",
            "line_spacing": 1.25,
            "hanging_indent_chars": 1.5,
            "space_before_pt": 0,
            "space_after_pt": 0,
        },
    }
    rules_path = tmp_path / "reference-rules.json"
    rules_path.write_text(
        json.dumps(rules_data, ensure_ascii=False), encoding="utf-8"
    )

    source = tmp_path / "reference-source.docx"
    document = Document()
    body = document.add_paragraph()
    body.add_run("已有研究")
    body.add_run("[1][2]")
    document.add_paragraph("参考文献")
    document.add_paragraph("[1] 第一条文献")
    document.add_paragraph("[2] Second reference")
    document.save(source)
    original_text = [paragraph.text for paragraph in Document(source).paragraphs]

    output = tmp_path / "reference-fixed.docx"
    result = fix_document(source, rules_path, output)

    fixed = Document(output)
    assert [paragraph.text for paragraph in fixed.paragraphs] == original_text
    assert result["summary"]["fixed_issues"] > 0
    assert {
        "_SDAU_REF_0001",
        "_SDAU_REF_0002",
        "_SDAU_CITE_0001_001",
        "_SDAU_CITE_0002_001",
    }.issubset(
        set(
            fixed.element.xpath(
                ".//*[local-name()='bookmarkStart']"
                "/@*[local-name()='name']"
            )
        )
    )
    final_report = inspect_document(output, rules_path)
    navigation_fields = {
        "citation_superscript",
        "citation_link",
        "citation_bookmark",
        "reference_bookmark",
        "reference_backlink",
    }
    assert not [
        issue
        for issue in final_report["issues"]
        if issue["field"] in navigation_fields
    ]
    for paragraph in fixed.paragraphs[2:]:
        assert paragraph.paragraph_format.line_spacing == pytest.approx(1.25)
        assert paragraph.paragraph_format.first_line_indent.cm < 0


def test_content_validation_allows_added_reference_navigation_but_preserves_existing(
    tmp_path: Path,
    rules_path: Path,
) -> None:
    from docx.oxml.ns import qn

    from word_format_tool.document_analyzer import analyze_document
    from word_format_tool.reference_links import (
        analyze_reference_document,
        apply_reference_navigation,
    )
    from word_format_tool.rule_loader import load_rules

    source = tmp_path / "navigation-source.docx"
    document = Document()
    existing = document.add_paragraph("既有书签")
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), "7")
    bookmark_start.set(qn("w:name"), "_ExistingBookmark")
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), "7")
    existing._p.insert(0, bookmark_start)
    existing._p.append(bookmark_end)
    document.add_paragraph("研究结果[1]")
    document.add_paragraph("参考文献")
    document.add_paragraph("[1] 第一条文献")
    document.save(source)

    candidate = tmp_path / "navigation-candidate.docx"
    loaded = Document(source)
    rules = load_rules(rules_path)
    reference_map = analyze_reference_document(
        loaded, analyze_document(loaded, rules)
    )
    apply_reference_navigation(loaded, reference_map)
    loaded.save(candidate)

    document_io.validate_preserved_content(source, candidate)

    damaged = tmp_path / "navigation-damaged.docx"
    damaged_document = Document(candidate)
    for element in damaged_document.element.xpath(
        ".//*[local-name()='bookmarkStart' "
        "and @*[local-name()='name']='_ExistingBookmark']"
    ):
        element.getparent().remove(element)
    damaged_document.save(damaged)

    with pytest.raises(DocumentReadError, match="navigation"):
        document_io.validate_preserved_content(source, damaged)

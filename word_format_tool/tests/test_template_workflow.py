from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from word_format_tool import build_template_rules, format_from_template


def test_template_rules_preserve_module_specific_formats(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    document = Document()

    abstract_heading = document.add_paragraph("摘 要")
    abstract_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    abstract_heading.runs[0].font.size = Pt(16)
    abstract_heading.runs[0].bold = True

    abstract_body = document.add_paragraph("模板摘要正文。")
    abstract_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract_body.runs[0].font.size = Pt(12)

    document.add_paragraph("关键词：模板；格式")

    toc_heading = document.add_paragraph("目 录")
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_heading.runs[0].font.size = Pt(16)
    toc_heading.runs[0].bold = True

    toc_entry = document.add_paragraph("1 绪论........................1")
    toc_entry.runs[0].font.size = Pt(10.5)

    document.add_paragraph("1 绪论")
    document.add_paragraph("模板正文。")

    reference_heading = document.add_paragraph("参考文献")
    reference_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    reference_heading.runs[0].font.size = Pt(16)
    reference_heading.runs[0].bold = True
    reference_entry = document.add_paragraph("[1] 模板文献")
    reference_entry.runs[0].font.size = Pt(10.5)
    document.save(template)

    rules = build_template_rules(template)

    assert rules["roles"]["abstract_heading_zh"]["font_size_pt"] == 16
    assert rules["roles"]["abstract_body_zh"]["font_size_pt"] == 12
    assert rules["roles"]["toc_heading"]["alignment"] == "center"
    assert rules["roles"]["toc_entry_1"]["font_size_pt"] == 10.5
    assert rules["roles"]["reference_heading"]["alignment"] == "center"
    assert rules["roles"]["reference_entry_zh"]["font_size_pt"] == 10.5


def test_format_from_template_runs_without_intermediate_rule_files(
    tmp_path: Path,
) -> None:
    template = tmp_path / "template.docx"
    template_document = Document()
    heading = template_document.add_paragraph("摘 要")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.size = Pt(16)
    heading.runs[0].bold = True
    body = template_document.add_paragraph("模板摘要正文。")
    body.runs[0].font.size = Pt(12)
    template_document.add_paragraph("关键词：模板")
    template_document.add_paragraph("1 绪论")
    template_document.add_paragraph("模板正文。")
    template_document.save(template)

    draft = tmp_path / "draft.docx"
    draft_document = Document()
    draft_heading = draft_document.add_paragraph("摘要")
    draft_heading.runs[0].font.size = Pt(10)
    draft_document.add_paragraph("初稿摘要正文。")
    draft_document.add_paragraph("关键词：初稿")
    draft_document.add_paragraph("1 绪论")
    draft_document.add_paragraph("初稿正文。")
    draft_document.save(draft)

    output = tmp_path / "formatted.docx"
    result = format_from_template(template, draft, output)

    formatted = Document(output)
    assert result["output"] == str(output)
    assert result["module_counts"]["abstract_heading_zh"] == 1
    assert formatted.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert formatted.paragraphs[0].runs[0].font.size.pt == 16

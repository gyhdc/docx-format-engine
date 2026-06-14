from __future__ import annotations

import importlib
import importlib.util
from pathlib import Path
from types import ModuleType

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from word_format_tool.document_analyzer import analyze_document
from word_format_tool.rule_loader import load_rules


def _reference_links() -> ModuleType:
    spec = importlib.util.find_spec("word_format_tool.reference_links")
    assert spec is not None, "reference_links module must exist"
    return importlib.import_module("word_format_tool.reference_links")


def test_citation_parser_supports_template_and_compound_markers() -> None:
    module = _reference_links()

    markers = module.find_citation_markers(
        "单条[1]，相邻[2][3]，并列[4,6，8]，区间[9-11]，"
        "混合［12,14–16］，非文献[abc]。"
    )

    assert [
        (marker.text, marker.numbers, marker.visible_numbers)
        for marker in markers
    ] == [
        ("[1]", (1,), (1,)),
        ("[2]", (2,), (2,)),
        ("[3]", (3,), (3,)),
        ("[4,6，8]", (4, 6, 8), (4, 6, 8)),
        ("[9-11]", (9, 10, 11), (9, 11)),
        ("［12,14–16］", (12, 14, 15, 16), (12, 14, 16)),
    ]


def test_reference_analysis_reports_duplicates_gaps_and_unresolved_citations(
    tmp_path: Path,
    rules_path: Path,
) -> None:
    module = _reference_links()
    source = tmp_path / "reference-analysis.docx"
    document = Document()
    document.add_paragraph("已有研究[1,2-4]。")
    document.add_paragraph("参考文献")
    document.add_paragraph("[1] 第一条")
    document.add_paragraph("[3] 第三条")
    document.add_paragraph("[3] 重复的第三条")
    document.add_paragraph("致谢")
    document.save(source)

    loaded = Document(source)
    analyzed = analyze_document(loaded, load_rules(rules_path))
    result = module.analyze_reference_document(loaded, analyzed)

    assert tuple(result.entries_by_number) == (1, 3)
    assert result.duplicate_entry_numbers == (3,)
    assert result.missing_entry_numbers == (2,)
    assert result.cited_numbers == (1, 2, 3, 4)
    assert result.unresolved_citation_numbers == (2, 4)
    assert result.uncited_entry_numbers == ()


def test_reference_analysis_reports_missing_leading_number(
    rules_path: Path,
) -> None:
    module = _reference_links()
    document = Document()
    document.add_paragraph("参考文献")
    document.add_paragraph("[2] 第二条文献")
    document.add_paragraph("[3] 第三条文献")
    rules = load_rules(rules_path)

    result = module.analyze_reference_document(
        document, analyze_document(document, rules)
    )

    assert result.missing_entry_numbers == (1,)


def test_navigation_adds_superscript_links_bookmarks_and_backlinks_idempotently(
    tmp_path: Path,
    rules_path: Path,
) -> None:
    module = _reference_links()
    source = tmp_path / "reference-links.docx"
    document = Document()
    body = document.add_paragraph()
    body.add_run("研究结果")
    citation = body.add_run("[1][2]")
    citation.italic = True
    document.add_paragraph("参考文献")
    document.add_paragraph("[1] 第一条文献")
    document.add_paragraph("[2] 第二条文献")
    document.save(source)

    loaded = Document(source)
    original_text = [paragraph.text for paragraph in loaded.paragraphs]
    analyzed = analyze_document(loaded, load_rules(rules_path))
    reference_map = module.analyze_reference_document(loaded, analyzed)

    first = module.apply_reference_navigation(loaded, reference_map)

    assert first.changed is True
    assert [paragraph.text for paragraph in loaded.paragraphs] == original_text
    bookmark_names = loaded.element.xpath(
        ".//*[local-name()='bookmarkStart']/@*[local-name()='name']"
    )
    assert {
        "_SDAU_REF_0001",
        "_SDAU_REF_0002",
        "_SDAU_CITE_0001_001",
        "_SDAU_CITE_0002_001",
    }.issubset(set(bookmark_names))

    body_anchors = loaded.paragraphs[0]._p.xpath(
        "./*[local-name()='hyperlink']/@*[local-name()='anchor']"
    )
    assert body_anchors == ["_SDAU_REF_0001", "_SDAU_REF_0002"]
    superscript_values = loaded.paragraphs[0]._p.xpath(
        "./*[local-name()='hyperlink']"
        "/*[local-name()='r']/*[local-name()='rPr']"
        "/*[local-name()='vertAlign']/@*[local-name()='val']"
    )
    assert superscript_values == ["superscript", "superscript"]
    italic_runs = loaded.paragraphs[0]._p.xpath(
        "./*[local-name()='hyperlink']"
        "/*[local-name()='r']/*[local-name()='rPr']/*[local-name()='i']"
    )
    assert len(italic_runs) == 2

    first_reference_anchor = loaded.paragraphs[2]._p.xpath(
        "./*[local-name()='hyperlink']/@*[local-name()='anchor']"
    )
    second_reference_anchor = loaded.paragraphs[3]._p.xpath(
        "./*[local-name()='hyperlink']/@*[local-name()='anchor']"
    )
    assert first_reference_anchor == ["_SDAU_CITE_0001_001"]
    assert second_reference_anchor == ["_SDAU_CITE_0002_001"]

    counts_before = (
        len(loaded.element.xpath(".//*[local-name()='bookmarkStart']")),
        len(loaded.element.xpath(".//*[local-name()='hyperlink']")),
    )
    refreshed = module.analyze_reference_document(
        loaded,
        analyze_document(loaded, load_rules(rules_path)),
    )
    second = module.apply_reference_navigation(loaded, refreshed)
    counts_after = (
        len(loaded.element.xpath(".//*[local-name()='bookmarkStart']")),
        len(loaded.element.xpath(".//*[local-name()='hyperlink']")),
    )

    assert second.changed is False
    assert counts_after == counts_before
    assert [paragraph.text for paragraph in loaded.paragraphs] == original_text


def test_navigation_links_citation_split_across_multiple_runs(
    tmp_path: Path,
    rules_path: Path,
) -> None:
    module = _reference_links()
    source = tmp_path / "split-reference-link.docx"
    document = Document()
    body = document.add_paragraph()
    body.add_run("已有研究")
    body.add_run("[").bold = True
    body.add_run("1").italic = True
    body.add_run("]").bold = True
    document.add_paragraph("参考文献")
    document.add_paragraph("[1] 第一条文献")
    document.save(source)

    loaded = Document(source)
    original_text = [paragraph.text for paragraph in loaded.paragraphs]
    rules = load_rules(rules_path)
    result = module.apply_reference_navigation(
        loaded,
        module.analyze_reference_document(
            loaded, analyze_document(loaded, rules)
        ),
    )

    assert result.changed is True
    assert [paragraph.text for paragraph in loaded.paragraphs] == original_text
    assert loaded.paragraphs[0]._p.xpath(
        "./*[local-name()='hyperlink' "
        "and @*[local-name()='anchor']='_SDAU_REF_0001']"
    )
    linked_number_runs = loaded.paragraphs[0]._p.xpath(
        "./*[local-name()='hyperlink' "
        "and @*[local-name()='anchor']='_SDAU_REF_0001']"
        "/*[local-name()='r']"
    )
    assert linked_number_runs[0].xpath(
        "./*[local-name()='rPr']/*[local-name()='i']"
    )
    marker_runs = loaded.paragraphs[0]._p.xpath(
        "./*[local-name()='r']"
        " | ./*[local-name()='hyperlink']/*[local-name()='r']"
    )
    assert all(
        run.xpath(
            "./*[local-name()='rPr']"
            "/*[local-name()='vertAlign' "
            "and @*[local-name()='val']='superscript']"
        )
        for run in marker_runs[-3:]
    )


def test_navigation_repairs_missing_backlink_when_reference_bookmark_exists(
    rules_path: Path,
) -> None:
    module = _reference_links()
    document = Document()
    document.add_paragraph("已有研究[1]。")
    document.add_paragraph("参考文献")
    entry = document.add_paragraph("[1] 第一条文献")
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), "9")
    bookmark_start.set(qn("w:name"), "_SDAU_REF_0001")
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), "9")
    entry._p.insert(0, bookmark_start)
    entry._p.append(bookmark_end)
    rules = load_rules(rules_path)

    result = module.apply_reference_navigation(
        document,
        module.analyze_reference_document(
            document, analyze_document(document, rules)
        ),
    )

    assert result.changed is True
    assert len(
        entry._p.xpath(
            "./*[local-name()='bookmarkStart' "
            "and @*[local-name()='name']='_SDAU_REF_0001']"
        )
    ) == 1
    assert entry._p.xpath(
        "./*[local-name()='hyperlink' "
        "and @*[local-name()='anchor']='_SDAU_CITE_0001_001']"
    )

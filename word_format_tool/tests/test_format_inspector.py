from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement

from word_format_tool.document_analyzer import analyze_document
from word_format_tool.format_inspector import inspect_analyzed_document
from word_format_tool.rule_loader import load_rules


def test_inspector_finds_body_and_heading_format_issues(
    draft_path: Path, rules_path: Path,
) -> None:
    rules = load_rules(rules_path)
    analyzed = analyze_document(Document(draft_path), rules)

    report = inspect_analyzed_document(
        Document(draft_path), analyzed, rules, draft_path, rules_path
    )

    fields_by_role = {
        (issue.role, issue.field)
        for issue in report.issues
    }
    assert ("heading_1", "font_size_pt") in fields_by_role
    assert ("heading_1", "keep_with_next") in fields_by_role
    assert ("body", "font_size_pt") in fields_by_role
    assert ("body", "alignment") in fields_by_role
    assert report.summary.total_issues > 0


def test_inspector_checks_page_margins(
    draft_path: Path, rules_path: Path,
) -> None:
    rules = load_rules(rules_path)
    document = Document(draft_path)
    analyzed = analyze_document(document, rules)

    report = inspect_analyzed_document(
        document, analyzed, rules, draft_path, rules_path
    )

    assert any(issue.field == "margin_left_cm" for issue in report.issues)


def test_inspector_accepts_correct_caption_adjacency(
    tmp_path: Path, rules_data: dict,
) -> None:
    rules_data["roles"]["figure_caption"] = {
        "expected_position": "below_object"
    }
    rules_data["roles"]["table_caption"] = {
        "expected_position": "above_object"
    }
    rules_path = tmp_path / "caption-rules.json"
    rules_path.write_text(
        json.dumps(rules_data, ensure_ascii=False), encoding="utf-8"
    )

    source = tmp_path / "captions.docx"
    document = Document()
    drawing_paragraph = document.add_paragraph()
    drawing_paragraph.add_run()._r.append(OxmlElement("w:drawing"))
    document.add_paragraph("图 1 正确位置")
    document.add_paragraph("表 1 正确位置")
    document.add_table(rows=1, cols=1)
    document.save(source)

    rules = load_rules(rules_path)
    loaded = Document(source)
    analyzed = analyze_document(loaded, rules)
    report = inspect_analyzed_document(
        loaded, analyzed, rules, source, rules_path
    )

    assert not [
        issue for issue in report.issues if issue.field == "expected_position"
    ]


def test_inspector_ignores_table_of_contents_without_unknown_warnings(
    tmp_path: Path, rules_path: Path,
) -> None:
    source = tmp_path / "toc.docx"
    document = Document()
    document.add_paragraph("目    录")
    toc_style = document.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
    entry = document.add_paragraph("1 引言\t1")
    entry.style = toc_style
    document.add_paragraph("1 引言")
    document.save(source)

    rules = load_rules(rules_path)
    loaded = Document(source)
    analyzed = analyze_document(loaded, rules)
    report = inspect_analyzed_document(
        loaded, analyzed, rules, source, rules_path
    )

    assert not [issue for issue in report.issues if issue.field == "role"]
    assert any("目录" in note and "2" in note for note in report.notes)


def test_inspector_reports_visible_table_of_contents_field_errors(
    tmp_path: Path, rules_path: Path,
) -> None:
    source = tmp_path / "broken-toc.docx"
    document = Document()
    toc_style = document.styles.add_style("TOC 2", WD_STYLE_TYPE.PARAGRAPH)
    entry = document.add_paragraph("3.2 抖音\t")
    entry.style = toc_style
    for fragment in (" PAGE", "REF _TocMissingBookmark \\h "):
        field_run = OxmlElement("w:r")
        instruction = OxmlElement("w:instrText")
        instruction.text = fragment
        field_run.append(instruction)
        entry._p.append(field_run)
    document.save(source)

    rules = load_rules(rules_path)
    loaded = Document(source)
    analyzed = analyze_document(loaded, rules)
    report = inspect_analyzed_document(
        loaded, analyzed, rules, source, rules_path
    )

    field_errors = [
        issue for issue in report.issues if issue.field == "field_result"
    ]
    assert len(field_errors) == 1
    assert field_errors[0].fixable is False
    assert "目录" in field_errors[0].message
    assert "_TocMissingBookmark" in field_errors[0].actual


def test_caption_adjacency_skips_blank_paragraphs(
    tmp_path: Path, rules_data: dict,
) -> None:
    rules_data["roles"]["figure_caption"] = {
        "expected_position": "below_object"
    }
    rules_data["roles"]["table_caption"] = {
        "expected_position": "above_object"
    }
    rules_path = tmp_path / "caption-rules.json"
    rules_path.write_text(
        json.dumps(rules_data, ensure_ascii=False), encoding="utf-8"
    )

    source = tmp_path / "captions-with-blanks.docx"
    document = Document()
    drawing_paragraph = document.add_paragraph()
    drawing_paragraph.add_run()._r.append(OxmlElement("w:drawing"))
    document.add_paragraph("")
    document.add_paragraph("图 1 空段落不应导致误报")
    document.add_paragraph("表 1 空段落不应导致误报")
    document.add_paragraph("")
    document.add_table(rows=1, cols=1)
    document.save(source)

    rules = load_rules(rules_path)
    loaded = Document(source)
    analyzed = analyze_document(loaded, rules)
    report = inspect_analyzed_document(
        loaded, analyzed, rules, source, rules_path
    )

    assert not [
        issue for issue in report.issues if issue.field == "expected_position"
    ]


def test_inspector_reports_reference_navigation_and_semantic_issues(
    tmp_path: Path,
    rules_path: Path,
) -> None:
    source = tmp_path / "reference-issues.docx"
    document = Document()
    document.add_paragraph("已有研究[1,2,4]。")
    document.add_paragraph("参考文献")
    document.add_paragraph("[1] 第一条文献")
    document.add_paragraph("[3] 第三条文献")
    document.add_paragraph("[3] 重复的第三条文献")
    document.save(source)

    rules = load_rules(rules_path)
    loaded = Document(source)
    analyzed = analyze_document(loaded, rules)
    report = inspect_analyzed_document(
        loaded, analyzed, rules, source, rules_path
    )

    fields = {issue.field for issue in report.issues}
    assert {
        "citation_superscript",
        "citation_link",
        "citation_bookmark",
        "reference_bookmark",
        "reference_backlink",
        "duplicate_reference_number",
        "missing_reference_number",
        "unresolved_citation",
        "uncited_reference",
    }.issubset(fields)
    semantic_fields = {
        "duplicate_reference_number",
        "missing_reference_number",
        "unresolved_citation",
        "uncited_reference",
    }
    assert all(
        issue.fixable is False
        for issue in report.issues
        if issue.field in semantic_fields
    )

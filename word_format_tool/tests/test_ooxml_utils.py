from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn

from word_format_tool.ooxml_utils import set_run_fonts


def test_set_run_fonts_sets_east_asia_ascii_and_hansi() -> None:
    run = Document().add_paragraph().add_run("中English")

    set_run_fonts(run, "SimSun", "Times New Roman")

    fonts = run._element.get_or_add_rPr().rFonts
    assert fonts.get(qn("w:eastAsia")) == "SimSun"
    assert fonts.get(qn("w:ascii")) == "Times New Roman"
    assert fonts.get(qn("w:hAnsi")) == "Times New Roman"

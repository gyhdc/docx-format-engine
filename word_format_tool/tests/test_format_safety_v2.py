from __future__ import annotations

import json

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt

from word_format_tool.document_analyzer import analyze_document
from word_format_tool.document_io import validate_preserved_content
from word_format_tool.exceptions import DocumentReadError
from word_format_tool.format_fixer import fix_document_to_path
from word_format_tool.format_inspector import inspect_analyzed_document
from word_format_tool.models import Rules
from word_format_tool.structure_tools import (
    complete_structure_to_path,
    plan_structure,
)
from word_format_tool.template_profiler import profile_template
from word_format_tool.visual_compare import RasterPage, compare_raster_pages


def _rules(
    *,
    roles: dict | None = None,
    page: dict | None = None,
    priority: dict | None = None,
) -> Rules:
    return Rules.model_validate(
        {
            "version": "0.1",
            "priority": priority or {},
            "page": page,
            "roles": roles or {"body": {}},
        }
    )


def test_contextual_title_roles_protect_other_front_matter() -> None:
    document = Document()
    school = document.add_paragraph("山东农业大学本科毕业论文")
    school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    school.runs[0].font.size = Pt(18)

    title_zh = document.add_paragraph("基于手势轨迹的可交互图像智能分析系统")
    title_zh.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_zh.runs[0].font.size = Pt(22)

    title_en = document.add_paragraph(
        "Development of an Interactive Image Analysis System Based on Gesture Trajectory"
    )
    title_en.runs[0].font.size = Pt(22)

    document.add_paragraph("目录")
    document.add_paragraph("1 引言")

    analyzed = analyze_document(
        document,
        _rules(
            roles={
                "title_zh": {},
                "title_en": {},
                "body": {},
                "heading_1": {},
            }
        ),
    )

    assert analyzed[0].role == "unknown"
    assert analyzed[0].ignore_reason == "protected_front_matter"
    assert analyzed[1].role == "title_zh"
    assert analyzed[2].role == "title_en"
    assert analyzed[3].ignore_reason == "table_of_contents"
    assert analyzed[4].role == "heading_1"


def test_legacy_title_style_is_split_by_language() -> None:
    document = Document()
    chinese = document.add_paragraph("中文论文题目")
    chinese.style = document.styles["Title"]
    english = document.add_paragraph("English Thesis Title")
    english.style = document.styles["Title"]

    analyzed = analyze_document(document, _rules())

    assert [item.role for item in analyzed] == ["title_zh", "title_en"]


def test_reference_entries_are_split_by_primary_language() -> None:
    document = Document()
    document.add_paragraph("参考文献")
    document.add_paragraph("[1] 张三. 智能分析系统研究[J]. 计算机学报, 2025.")
    document.add_paragraph(
        "[2] Smith J. Interactive image analysis[J]. Pattern Recognition, 2025."
    )

    analyzed = analyze_document(document, _rules())

    assert [item.role for item in analyzed] == [
        "reference_heading",
        "reference_entry_zh",
        "reference_entry_en",
    ]


def test_analyzer_inventories_cover_table_header_and_footer() -> None:
    document = Document()
    document.add_paragraph("论文题目")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "题目"
    table.cell(0, 1).text = "基于手势轨迹的智能分析系统"
    document.add_paragraph("1 引言")
    document.sections[0].header.paragraphs[0].text = "山东农业大学本科毕业论文"
    document.sections[0].footer.paragraphs[0].text = "第 1 页"

    analyzed = analyze_document(
        document,
        _rules(
            roles={
                "body": {},
                "cover_label": {},
                "cover_value": {},
                "header": {},
                "footer": {},
            }
        ),
    )

    scoped = {(item.story, item.role, item.text) for item in analyzed}
    assert ("table_cell", "cover_label", "题目") in scoped
    assert ("table_cell", "cover_value", "基于手势轨迹的智能分析系统") in scoped
    assert ("header", "header", "山东农业大学本科毕业论文") in scoped
    assert ("footer", "footer", "第 1 页") in scoped
    assert all(item.location for item in analyzed)


def test_first_body_table_is_not_misclassified_as_cover() -> None:
    document = Document()
    document.add_heading("1 引言", level=1)
    document.add_paragraph("正文")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "实验数据"

    analyzed = analyze_document(document, _rules())

    table_item = next(item for item in analyzed if item.story == "table_cell")
    assert table_item.role == "table_text"


def test_table_cell_inventory_is_stable_across_repeated_analysis() -> None:
    document = Document()
    table = document.add_table(rows=3, cols=3)
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            cell.text = f"{row_index}-{cell_index}"
    rules = _rules()

    counts = [
        sum(
            item.story == "table_cell"
            for item in analyze_document(document, rules)
        )
        for _ in range(10)
    ]

    assert counts == [9] * 10


def test_cover_table_blocks_page_geometry_repairs_by_default(tmp_path) -> None:
    document = Document()
    document.add_table(rows=1, cols=1).cell(0, 0).text = "封面"
    document.add_paragraph("正文")
    document.sections[0].left_margin = Cm(1)
    rules = _rules(page={"margin_left_cm": 3.0})
    analyzed = analyze_document(document, rules)

    report = inspect_analyzed_document(
        document, analyzed, rules, tmp_path / "draft.docx", tmp_path / "rules.json"
    )

    margin_issue = next(issue for issue in report.issues if issue.field == "margin_left_cm")
    assert margin_issue.fixable is False
    assert any(issue.field == "page_layout_risk" for issue in report.issues)
    assert report.coverage.uncovered_areas


def test_explicit_unsafe_page_geometry_override_is_honored(tmp_path) -> None:
    document = Document()
    document.add_table(rows=1, cols=1).cell(0, 0).text = "封面"
    document.add_paragraph("正文")
    document.sections[0].left_margin = Cm(1)
    rules = _rules(
        page={"margin_left_cm": 3.0},
        priority={"allow_unsafe_page_geometry": True},
    )
    analyzed = analyze_document(document, rules)

    report = inspect_analyzed_document(
        document, analyzed, rules, tmp_path / "draft.docx", tmp_path / "rules.json"
    )

    margin_issue = next(issue for issue in report.issues if issue.field == "margin_left_cm")
    assert margin_issue.fixable is True


def test_floating_anchor_is_reported_as_uncovered_layout_area(tmp_path) -> None:
    document = Document()
    paragraph = document.add_paragraph("正文")
    drawing = OxmlElement("w:drawing")
    drawing.append(OxmlElement("wp:anchor"))
    paragraph.runs[0]._r.append(drawing)
    rules = _rules()
    analyzed = analyze_document(document, rules)

    report = inspect_analyzed_document(
        document, analyzed, rules, tmp_path / "draft.docx", tmp_path / "rules.json"
    )

    assert report.coverage.layout_risk["floating_anchor_count"] == 1
    assert any("floating" in area for area in report.coverage.uncovered_areas)


def test_header_floating_anchor_is_included_in_layout_risk(tmp_path) -> None:
    document = Document()
    header = document.sections[0].header.paragraphs[0]
    header.text = "页眉"
    drawing = OxmlElement("w:drawing")
    drawing.append(OxmlElement("wp:anchor"))
    header.runs[0]._r.append(drawing)
    rules = _rules()
    analyzed = analyze_document(document, rules)

    report = inspect_analyzed_document(
        document, analyzed, rules, tmp_path / "draft.docx", tmp_path / "rules.json"
    )

    assert report.coverage.layout_risk["floating_anchor_count"] == 1


def test_fix_applies_explicit_cover_header_and_footer_rules(tmp_path) -> None:
    source = tmp_path / "scoped.docx"
    output = tmp_path / "scoped-fixed.docx"
    rules_path = tmp_path / "rules.json"
    document = Document()
    document.add_paragraph("正文")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "题目"
    table.cell(0, 1).text = "智能分析系统"
    document.add_paragraph("1 引言")
    document.sections[0].header.paragraphs[0].text = "学校页眉"
    document.sections[0].footer.paragraphs[0].text = "第 1 页"
    document.save(source)
    rules_path.write_text(
        json.dumps(
            {
                "version": "0.1",
                "roles": {
                    "body": {},
                    "cover_label": {"font_ascii": "Arial", "bold": True},
                    "cover_value": {
                        "font_ascii": "Times New Roman",
                        "font_size_pt": 16,
                    },
                    "header": {"font_ascii": "Calibri", "alignment": "center"},
                    "footer": {"font_ascii": "Courier New", "alignment": "center"},
                },
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    report = fix_document_to_path(source, rules_path, output)
    fixed = Document(output)

    assert fixed.tables[0].cell(0, 0).paragraphs[0].runs[0].font.name == "Arial"
    assert fixed.tables[0].cell(0, 0).paragraphs[0].runs[0].bold is True
    assert (
        fixed.tables[0].cell(0, 1).paragraphs[0].runs[0].font.name
        == "Times New Roman"
    )
    assert fixed.sections[0].header.paragraphs[0].runs[0].font.name == "Calibri"
    assert fixed.sections[0].footer.paragraphs[0].runs[0].font.name == "Courier New"
    assert report.summary.fixed_issues > 0


def test_fix_applies_different_chinese_and_english_reference_alignment(
    tmp_path,
) -> None:
    source = tmp_path / "references.docx"
    output = tmp_path / "references-fixed.docx"
    rules_path = tmp_path / "rules.json"
    document = Document()
    document.add_paragraph("参考文献")
    chinese = document.add_paragraph("[1] 张三. 智能分析研究[J]. 计算机学报, 2025.")
    english = document.add_paragraph(
        "[2] Smith J. Interactive analysis[J]. Pattern Recognition, 2025."
    )
    chinese.alignment = WD_ALIGN_PARAGRAPH.LEFT
    english.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    document.save(source)
    rules_path.write_text(
        json.dumps(
            {
                "version": "0.1",
                "roles": {
                    "reference_entry_zh": {"alignment": "both"},
                    "reference_entry_en": {"alignment": "left"},
                },
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    fix_document_to_path(source, rules_path, output)
    fixed = Document(output)

    assert fixed.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert fixed.paragraphs[2].alignment == WD_ALIGN_PARAGRAPH.LEFT


def test_fix_repairs_unambiguous_toc_bookmark_and_requests_field_update(
    tmp_path,
) -> None:
    source = tmp_path / "toc.docx"
    output = tmp_path / "toc-fixed.docx"
    rules_path = tmp_path / "rules.json"
    document = Document()
    toc_style = document.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
    toc = document.add_paragraph(style=toc_style)
    toc.add_run("1 引言\t")
    field_run = toc.add_run()
    instruction = OxmlElement("w:instrText")
    instruction.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instruction.text = " PAGEREF missing_heading \\h "
    field_run._r.append(instruction)
    toc.add_run("1")
    document.add_heading("1 引言", level=1)
    document.add_paragraph("正文。")
    document.save(source)
    rules_path.write_text(
        json.dumps(
            {
                "version": "0.1",
                "roles": {
                    "table_of_contents": {"font_size_pt": 10.5},
                    "heading_1": {},
                    "body": {},
                },
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    report = fix_document_to_path(source, rules_path, output)
    fixed = Document(output)
    field_code = "".join(
        fixed.paragraphs[0]._p.xpath(
            ".//*[local-name()='instrText']/text()"
        )
    )
    bookmark_names = fixed.element.xpath(
        ".//*[local-name()='bookmarkStart']/@*[local-name()='name']"
    )
    update_fields = fixed.settings.element.xpath(
        "./*[local-name()='updateFields']/@*[local-name()='val']"
    )

    assert "missing_heading" not in field_code
    assert any(name.startswith("_SDAU_TOC_") for name in bookmark_names)
    assert update_fields == ["true"]
    assert fixed.paragraphs[0].runs[0].font.size.pt == 10.5
    assert any(
        issue.field == "field_result" and issue.fixed for issue in report.issues
    )


def test_partial_toc_repair_does_not_request_global_field_refresh(
    tmp_path,
) -> None:
    source = tmp_path / "partial-toc.docx"
    output = tmp_path / "partial-toc-fixed.docx"
    rules_path = tmp_path / "rules.json"
    document = Document()
    toc_style = document.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
    for visible_text, target in (
        ("1 引言\t1", "missing_intro"),
        ("2 不存在章节\t2", "missing_unknown"),
    ):
        toc = document.add_paragraph(style=toc_style)
        toc.add_run(visible_text)
        field_run = toc.add_run()
        instruction = OxmlElement("w:instrText")
        instruction.text = f" PAGEREF {target} \\h "
        field_run._r.append(instruction)
    document.add_heading("1 引言", level=1)
    document.save(source)
    rules_path.write_text(
        json.dumps(
            {
                "version": "0.1",
                "roles": {
                    "table_of_contents": {},
                    "heading_1": {},
                },
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    fix_document_to_path(source, rules_path, output)
    fixed = Document(output)
    field_codes = [
        "".join(
            paragraph._p.xpath(
                ".//*[local-name()='instrText']/text()"
            )
        )
        for paragraph in fixed.paragraphs[:2]
    ]

    assert "_SDAU_TOC_" in field_codes[0]
    assert "missing_unknown" in field_codes[1]
    assert not fixed.settings.element.xpath(
        "./*[local-name()='updateFields']"
    )


def test_template_profile_and_structure_plan_expose_missing_sections(
    tmp_path,
) -> None:
    template_path = tmp_path / "template.docx"
    draft_path = tmp_path / "draft.docx"
    template = Document()
    template.add_paragraph("摘要：模板摘要")
    template.add_paragraph("关键词：智能；分析")
    template.add_heading("1 引言", level=1)
    template.add_paragraph("参考文献")
    template.save(template_path)
    draft = Document()
    draft.add_heading("1 引言", level=1)
    draft.add_paragraph("正文")
    draft.add_paragraph("参考文献")
    draft.save(draft_path)

    profile = profile_template(template_path)
    plan = plan_structure(template_path, draft_path)

    assert [item["role"] for item in profile["structure_outline"][:3]] == [
        "abstract",
        "keywords_zh",
        "heading_1",
    ]
    assert plan["missing_roles"] == ["abstract", "keywords"]
    assert "arbitrary_template_blocks" in plan["not_automatically_copied"]


def test_explicit_structure_completion_inserts_only_declared_missing_section(
    tmp_path,
) -> None:
    source = tmp_path / "structure.docx"
    output = tmp_path / "structure-completed.docx"
    rules_path = tmp_path / "rules.json"
    document = Document()
    document.add_heading("1 引言", level=1)
    document.add_paragraph("原始正文")
    document.add_paragraph("参考文献")
    document.save(source)
    rules_path.write_text(
        json.dumps(
            {
                "version": "0.1",
                "roles": {
                    "acknowledgements": {
                        "font_ascii": "Times New Roman",
                        "alignment": "center",
                    },
                    "body": {},
                },
                "structure": {
                    "required_sections": [
                        {
                            "role": "acknowledgements",
                            "heading_text": "致谢",
                            "placeholder_text": "请在此补充致谢内容。",
                            "insert_before_role": "reference_heading",
                        }
                    ]
                },
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    result = complete_structure_to_path(source, rules_path, output)
    completed = Document(output)
    texts = [paragraph.text for paragraph in completed.paragraphs]

    assert texts == [
        "1 引言",
        "原始正文",
        "致谢",
        "请在此补充致谢内容。",
        "参考文献",
    ]
    assert result["inserted_roles"] == ["acknowledgements"]
    assert Document(source).paragraphs[1].text == "原始正文"


def test_content_validation_detects_floating_object_position_change(
    tmp_path,
) -> None:
    source = tmp_path / "floating-source.docx"
    candidate = tmp_path / "floating-candidate.docx"
    document = Document()
    paragraph = document.add_paragraph("正文")
    drawing = OxmlElement("w:drawing")
    anchor = OxmlElement("wp:anchor")
    position = OxmlElement("wp:positionH")
    position.set("relativeFrom", "page")
    offset = OxmlElement("wp:posOffset")
    offset.text = "1000"
    position.append(offset)
    anchor.append(position)
    drawing.append(anchor)
    paragraph.runs[0]._r.append(drawing)
    document.save(source)
    changed = Document(source)
    changed.element.xpath(".//*[local-name()='posOffset']")[0].text = "9000"
    changed.save(candidate)

    with pytest.raises(DocumentReadError, match="layout_objects"):
        validate_preserved_content(source, candidate)


def test_content_validation_detects_header_text_change(tmp_path) -> None:
    source = tmp_path / "header-source.docx"
    candidate = tmp_path / "header-candidate.docx"
    document = Document()
    document.sections[0].header.paragraphs[0].text = "原页眉"
    document.save(source)
    changed = Document(source)
    changed.sections[0].header.paragraphs[0].text = "被改坏的页眉"
    changed.save(candidate)

    with pytest.raises(DocumentReadError, match="header_footer_text"):
        validate_preserved_content(source, candidate)


def test_raster_visual_compare_reports_page_size_and_pixel_changes() -> None:
    white = RasterPage(width=2, height=2, pixels=bytes([255] * 12))
    changed = RasterPage(
        width=2,
        height=2,
        pixels=bytes([0, 0, 0] + [255] * 9),
    )
    wider = RasterPage(width=3, height=2, pixels=bytes([255] * 18))

    pixel_report = compare_raster_pages([white], [changed])
    size_report = compare_raster_pages([white], [wider])
    count_report = compare_raster_pages([white], [white, white])

    assert pixel_report["pages"][0]["changed_pixel_ratio"] == 0.25
    assert pixel_report["status"] == "warning"
    assert size_report["page_size_changed"] is True
    assert count_report["page_count_changed"] is True

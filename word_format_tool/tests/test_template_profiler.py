from __future__ import annotations

from pathlib import Path

from docx import Document

from word_format_tool.template_profiler import extract_samples, profile_template


def test_profile_template_extracts_page_styles_and_samples(tmp_path: Path) -> None:
    path = tmp_path / "template.docx"
    document = Document()
    document.add_heading("1 引言", level=1)
    document.add_paragraph("模板正文样本")
    document.save(path)

    profile = profile_template(path)

    assert profile["source_file"] == "template.docx"
    assert profile["page"]["sections"]
    assert any(style["name"] == "Heading 1" for style in profile["styles"])
    assert profile["paragraph_samples"][0]["guessed_role"] == "heading_1"


def test_extract_samples_returns_representative_paragraphs(tmp_path: Path) -> None:
    path = tmp_path / "template.docx"
    document = Document()
    document.add_paragraph("摘要：测试")
    document.add_paragraph("图 1 测试结果")
    document.save(path)

    samples = extract_samples(path)

    assert [sample["guessed_role"] for sample in samples] == [
        "abstract",
        "figure_caption",
    ]

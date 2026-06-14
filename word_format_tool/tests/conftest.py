from __future__ import annotations

import json
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


@pytest.fixture
def rules_data() -> dict[str, Any]:
    return {
        "version": "0.1",
        "priority": {
            "explicit_rules_over_template": True,
            "preserve_content": True,
            "preserve_images": True,
            "preserve_tables": True,
        },
        "page": {
            "paper_size": "A4",
            "margin_top_cm": 2.5,
            "margin_bottom_cm": 2.5,
            "margin_left_cm": 3.0,
            "margin_right_cm": 2.5,
        },
        "roles": {
            "title": {
                "font_east_asia": "SimHei",
                "font_ascii": "Times New Roman",
                "font_size_pt": 16,
                "bold": True,
                "alignment": "center",
            },
            "heading_1": {
                "font_east_asia": "SimHei",
                "font_ascii": "Times New Roman",
                "font_size_pt": 15,
                "bold": True,
                "alignment": "left",
                "keep_with_next": True,
            },
            "body": {
                "font_east_asia": "SimSun",
                "font_ascii": "Times New Roman",
                "font_size_pt": 12,
                "bold": False,
                "line_spacing": 1.5,
                "first_line_indent_chars": 2,
                "alignment": "both",
                "space_before_pt": 0,
                "space_after_pt": 0,
            },
        },
        "detection": {
            "heading_1_patterns": [r"^\d+\s+"],
            "heading_2_patterns": [r"^\d+\.\d+\s+"],
            "heading_3_patterns": [r"^\d+\.\d+\.\d+\s+"],
            "figure_caption_patterns": [r"^图\s*\d+"],
            "table_caption_patterns": [r"^表\s*\d+"],
            "reference_patterns": [r"^\[\d+\]"],
        },
    }


@pytest.fixture
def rules_path(tmp_path: Path, rules_data: dict[str, Any]) -> Path:
    path = tmp_path / "rules.json"
    path.write_text(json.dumps(rules_data, ensure_ascii=False), encoding="utf-8")
    return path


@pytest.fixture
def draft_path(tmp_path: Path) -> Path:
    path = tmp_path / "draft.docx"
    document = Document()

    title = document.add_paragraph("论文题目")
    title.style = document.styles["Title"]

    heading = document.add_paragraph("1 引言")
    heading.style = document.styles["Normal"]
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading.runs[0]
    heading_run.font.size = Pt(10)
    heading_run.bold = False

    body = document.add_paragraph("这是一段用于测试格式检查和修复的正文。")
    body_run = body.runs[0]
    body_run.font.name = "Arial"
    body_run.font.size = Pt(10)
    body_run.bold = True
    body.alignment = WD_ALIGN_PARAGRAPH.LEFT

    unknown = document.add_paragraph("")
    unknown.style = document.styles["Normal"]
    document.save(path)
    return path

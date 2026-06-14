from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from word_format_tool import complete_structure, plan_structure, profile_template
from word_format_tool.models import (
    FormatReport,
    StructureCompletionResult,
    StructurePlan,
    TemplateProfile,
    VisualComparisonReport,
)

PACKAGE_ROOT = Path(__file__).resolve().parents[1]


def test_public_response_models_publish_json_schema() -> None:
    expected_required_fields = {
        TemplateProfile: {"source_file", "page", "paragraph_samples", "layout_risk"},
        StructurePlan: {"template", "draft", "missing_roles"},
        StructureCompletionResult: {"source", "output", "inserted_roles"},
        VisualComparisonReport: {
            "status",
            "source_page_count",
            "candidate_page_count",
            "pages",
        },
        FormatReport: {"document", "rules", "phase", "summary", "issues"},
    }

    for model, required_fields in expected_required_fields.items():
        schema = model.model_json_schema()
        assert required_fields <= set(schema["required"])


def test_public_structured_results_match_their_response_models(
    tmp_path: Path,
) -> None:
    template_path = tmp_path / "template.docx"
    draft_path = tmp_path / "draft.docx"
    rules_path = tmp_path / "rules.json"
    completed_path = tmp_path / "completed.docx"

    template = Document()
    template.add_paragraph("摘要：模板摘要")
    template.add_heading("1 引言", level=1)
    template.add_paragraph("参考文献")
    template.save(template_path)

    draft = Document()
    draft.add_heading("1 引言", level=1)
    draft.add_paragraph("正文")
    draft.add_paragraph("参考文献")
    draft.save(draft_path)

    rules_path.write_text(
        json.dumps(
            {
                "version": "0.1",
                "roles": {
                    "body": {},
                    "acknowledgements": {"alignment": "center"},
                },
                "structure": {
                    "required_sections": [
                        {
                            "role": "acknowledgements",
                            "heading_text": "致谢",
                            "insert_before_role": "reference_heading",
                        }
                    ]
                },
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    TemplateProfile.model_validate(profile_template(template_path))
    StructurePlan.model_validate(plan_structure(template_path, draft_path))
    result = complete_structure(draft_path, rules_path, completed_path)
    StructureCompletionResult.model_validate(result)
    assert result["inserted_roles"] == ["acknowledgements"]


def test_agent_integration_doc_covers_every_public_api_and_contract() -> None:
    integration_doc = PACKAGE_ROOT / "AGENT_INTEGRATION.md"
    readme = PACKAGE_ROOT / "README.md"

    assert integration_doc.is_file()
    text = integration_doc.read_text(encoding="utf-8")
    for function_name in (
        "profile_template",
        "validate_rules",
        "inspect_document",
        "inspect_structure",
        "fix_document",
        "apply_structure_operations",
        "validate_layout",
        "refresh_fields",
        "link_references",
        "plan_structure",
        "complete_structure",
        "compare_visual",
        "rules_schema",
    ):
        assert f"`{function_name}" in text

    for model_name in (
        "TemplateProfile",
        "Rules",
        "FormatReport",
        "StructurePlan",
        "StructureCompletionResult",
        "VisualComparisonReport",
        "ParagraphSelector",
        "StructureOperationPlan",
        "LayoutExpectations",
    ):
        assert f"`{model_name}`" in text

    for heading in (
        "## 能力边界总览",
        "## 按能力调用",
        "## profile → rules 映射",
        "## API 返回契约",
        "## 错误合约",
        "## Agent 端到端流程",
    ):
        assert heading in text

    capability_sections = {
        "### 1. 提取模板事实": ("profile_template", "profile-template"),
        "### 2. 校验规则": ("validate_rules", "validate-rules"),
        "### 3. 检查与预演": ("inspect_document", "dry-run"),
        "### 4. 执行格式修复": ("fix_document", "wordfmt fix"),
        "### 5. 仅处理文献跳转": ("link_references", "link-references"),
        "### 6. 规划文档结构": ("plan_structure", "plan-structure"),
        "### 7. 补建缺失章节": ("complete_structure", "complete-structure"),
        "### 8. 视觉回归检查": ("compare_visual", "visual-compare"),
    }
    for heading, required_terms in capability_sections.items():
        assert heading in text
        section = text.split(heading, maxsplit=1)[1].split("\n### ", maxsplit=1)[0]
        for term in (*required_terms, "能做", "不能做", "验收"):
            assert term in section

    for boundary in (
        "不自动生成可靠的 `rules.json`",
        "不自动复制模板正文",
        "不编辑文本框或浮动对象",
        "不覆盖原始输入文件",
        "不猜测无法唯一匹配",
    ):
        assert boundary in text

    assert "AGENT_INTEGRATION.md" in readme.read_text(encoding="utf-8")

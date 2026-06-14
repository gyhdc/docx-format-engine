from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typer.testing import CliRunner

import word_format_tool as public_api
from word_format_tool.cli import app
from word_format_tool.exceptions import (
    FieldRefreshIntegrityError,
    StructureOperationError,
)
from word_format_tool.layout_validator import validate_layout
from word_format_tool.models import ParagraphSelector
from word_format_tool.optional_word_backend import refresh_fields_to_path
from word_format_tool.paragraph_locator import locate_paragraph
from word_format_tool.structure_inspector import inspect_structure
from word_format_tool.structure_operations import apply_structure_operations_to_path
from word_format_tool.template_rules import build_template_rules


runner = CliRunner()


def _document_with_toc_and_body(path: Path) -> Path:
    document = Document()
    document.add_paragraph("封面")
    document.add_paragraph("目录")
    toc_style = document.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
    document.add_paragraph("1 任务描述\t3", style=toc_style)
    document.add_paragraph("2 方案设计\t4", style=toc_style)
    document.add_heading("1 任务描述", level=1)
    document.add_paragraph("正文第一段。")
    document.add_heading("2 方案设计", level=1)
    document.add_paragraph("正文第二段。")
    document.save(path)
    return path


def _body_selector() -> dict[str, object]:
    return {
        "text": "1 任务描述",
        "role": "heading_1",
        "after_role": "table_of_contents",
    }


def _add_outer_toc_field_ending_in_body_heading(path: Path) -> Path:
    document = Document(path)
    toc_entry = document.paragraphs[2]
    toc_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    toc_run.extend((begin, instruction, separate))
    toc_entry._p.insert(1, toc_run)

    body_heading = document.paragraphs[4]
    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    body_heading._p.insert(1, end_run)
    document.save(path)
    return path


def _operation_plan(path: Path) -> Path:
    path.write_text(
        json.dumps(
            {
                "version": "0.1",
                "operations": [
                    {
                        "type": "insert_section_before",
                        "target": _body_selector(),
                        "break_type": "next_page",
                    },
                    {
                        "type": "set_header",
                        "section_start": _body_selector(),
                        "text": "山东农业大学学士学位论文",
                        "alignment": "center",
                        "bottom_border": True,
                    },
                    {
                        "type": "set_page_number",
                        "section_start": _body_selector(),
                        "start": 1,
                        "alignment": "center",
                    },
                    {"type": "request_field_update"},
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    return path


def _expectations(path: Path) -> Path:
    path.write_text(
        json.dumps(
            {
                "version": "0.1",
                "section_count_at_least": 2,
                "body_start": _body_selector(),
                "body_first_heading_equals": "1 任务描述",
                "toc_before_body": True,
                "section_header_equals": "山东农业大学学士学位论文",
                "section_has_page_field": True,
                "section_page_number_starts_at": 1,
                "update_fields_enabled": True,
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    return path


def test_locator_uses_toc_boundary_to_find_real_body_heading(tmp_path: Path) -> None:
    source = _document_with_toc_and_body(tmp_path / "source.docx")
    document = Document(source)

    located = locate_paragraph(
        document,
        ParagraphSelector.model_validate(_body_selector()),
    )

    assert located.paragraph_index == 4
    assert located.text == "1 任务描述"
    assert located.role == "heading_1"


def test_locator_rejects_ambiguous_match_without_occurrence(tmp_path: Path) -> None:
    source = tmp_path / "ambiguous.docx"
    document = Document()
    document.add_paragraph("重复内容")
    document.add_paragraph("重复内容")
    document.save(source)

    with pytest.raises(StructureOperationError, match="匹配到 2 个"):
        locate_paragraph(
            Document(source),
            ParagraphSelector(text="重复内容"),
        )


def test_structure_inspection_reports_sections_fields_and_toc_boundary(
    tmp_path: Path,
) -> None:
    source = _document_with_toc_and_body(tmp_path / "source.docx")

    report = inspect_structure(source)

    assert report["section_count"] == 1
    assert report["toc"]["start_paragraph"] == 1
    assert report["toc"]["end_paragraph"] == 3
    body_heading = next(
        item
        for item in report["paragraphs"]
        if item["text"] == "1 任务描述" and item["role"] == "heading_1"
    )
    assert body_heading["section_index"] == 0
    assert body_heading["starts_section"] is False
    assert body_heading["locator"]["after_role"] == "table_of_contents"


def test_operations_create_idempotent_body_section_header_and_page_number(
    tmp_path: Path,
) -> None:
    source = _document_with_toc_and_body(tmp_path / "source.docx")
    plan = _operation_plan(tmp_path / "operations.json")
    first_output = tmp_path / "first.docx"
    second_output = tmp_path / "second.docx"

    first_report = apply_structure_operations_to_path(source, plan, first_output)
    second_report = apply_structure_operations_to_path(
        first_output, plan, second_output
    )
    first = inspect_structure(first_output)
    second = inspect_structure(second_output)

    assert first_report["changed"] is True
    assert first["section_count"] == 2
    assert first["sections"][0]["header_text"] == ""
    assert first["sections"][1]["header_text"] == "山东农业大学学士学位论文"
    assert "PAGE" in first["sections"][1]["footer_field_codes"]
    assert first["sections"][1]["page_number_start"] == 1
    assert first["update_fields_enabled"] is True

    assert second_report["changed"] is False
    assert second["section_count"] == 2
    assert second["sections"][1]["footer_field_codes"].count("PAGE") == 1


def test_existing_empty_heading_section_carrier_is_normalized(
    tmp_path: Path,
) -> None:
    source = _document_with_toc_and_body(tmp_path / "source.docx")
    plan = _operation_plan(tmp_path / "operations.json")
    first_output = tmp_path / "first.docx"
    vulnerable = tmp_path / "vulnerable.docx"
    normalized = tmp_path / "normalized.docx"
    unchanged = tmp_path / "unchanged.docx"
    apply_structure_operations_to_path(source, plan, first_output)

    document = Document(first_output)
    heading_index = next(
        index
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text == "1 任务描述"
    )
    carrier = document.paragraphs[heading_index - 1]
    style = OxmlElement("w:pStyle")
    style.set(qn("w:val"), "Heading2")
    carrier._p.get_or_add_pPr().insert(0, style)
    document.save(vulnerable)

    first_report = apply_structure_operations_to_path(vulnerable, plan, normalized)
    second_report = apply_structure_operations_to_path(normalized, plan, unchanged)
    normalized_document = Document(normalized)
    normalized_heading_index = next(
        index
        for index, paragraph in enumerate(normalized_document.paragraphs)
        if paragraph.text == "1 任务描述"
    )
    normalized_carrier = normalized_document.paragraphs[normalized_heading_index - 1]

    assert first_report["operations"][0]["changed"] is True
    assert not normalized_carrier._p.xpath(
        "./*[local-name()='pPr']/*[local-name()='pStyle']"
    )
    assert "WORDFMT_SECTION_BREAK" in normalized_carrier.text
    assert normalized_carrier.runs[-1].font.hidden is True
    assert second_report["changed"] is False


def test_section_operation_moves_outer_toc_end_before_break_carrier(
    tmp_path: Path,
) -> None:
    source = _add_outer_toc_field_ending_in_body_heading(
        _document_with_toc_and_body(tmp_path / "source.docx")
    )
    plan = _operation_plan(tmp_path / "operations.json")
    output = tmp_path / "output.docx"
    unchanged = tmp_path / "unchanged.docx"

    first_report = apply_structure_operations_to_path(source, plan, output)
    second_report = apply_structure_operations_to_path(output, plan, unchanged)
    document = Document(output)
    heading_index = next(
        index
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text == "1 任务描述"
    )
    carrier = document.paragraphs[heading_index - 1]
    toc_tail = document.paragraphs[heading_index - 2]

    assert not document.paragraphs[heading_index]._p.xpath(
        ".//*[local-name()='fldChar' and @*[local-name()='fldCharType']='end']"
    )
    assert toc_tail._p.xpath(
        ".//*[local-name()='fldChar' and @*[local-name()='fldCharType']='end']"
    )
    assert carrier._p.xpath(
        "./*[local-name()='pPr']/*[local-name()='sectPr']"
    )
    assert first_report["operations"][0]["moved_toc_field_end"] is True
    assert second_report["changed"] is False


def test_failed_operation_does_not_publish_output(tmp_path: Path) -> None:
    source = _document_with_toc_and_body(tmp_path / "source.docx")
    output = tmp_path / "output.docx"
    plan = tmp_path / "bad-plan.json"
    plan.write_text(
        json.dumps(
            {
                "version": "0.1",
                "operations": [
                    {
                        "type": "insert_section_before",
                        "target": {"text": "不存在的标题"},
                        "break_type": "next_page",
                    }
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    with pytest.raises(StructureOperationError, match="未找到"):
        apply_structure_operations_to_path(source, plan, output)

    assert not output.exists()


def test_layout_validation_passes_machine_readable_structure_checks(
    tmp_path: Path,
) -> None:
    source = _document_with_toc_and_body(tmp_path / "source.docx")
    plan = _operation_plan(tmp_path / "operations.json")
    output = tmp_path / "output.docx"
    apply_structure_operations_to_path(source, plan, output)

    report = validate_layout(output, _expectations(tmp_path / "expectations.json"))

    assert report["status"] == "passed"
    assert all(check["status"] == "passed" for check in report["checks"])


def test_layout_validation_explains_missing_body_section(tmp_path: Path) -> None:
    source = _document_with_toc_and_body(tmp_path / "source.docx")

    report = validate_layout(source, _expectations(tmp_path / "expectations.json"))

    assert report["status"] == "failed"
    failed_codes = {
        check["code"] for check in report["checks"] if check["status"] == "failed"
    }
    assert "body_starts_new_section" in failed_codes
    assert "section_header_equals" in failed_codes
    assert "section_has_page_field" in failed_codes


def test_template_rule_extraction_drops_absurd_multiple_line_spacing(
    tmp_path: Path,
) -> None:
    template_path = tmp_path / "template.docx"
    document = Document()
    paragraph = document.add_paragraph("正文")
    properties = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "241300")
    spacing.set(qn("w:lineRule"), "auto")
    properties.append(spacing)
    document.save(template_path)

    rules = build_template_rules(template_path, modules=["body"])

    assert "line_spacing" not in rules["roles"]["body"]


def test_footer_text_and_paragraph_pagination_operations_are_idempotent(
    tmp_path: Path,
) -> None:
    source = _document_with_toc_and_body(tmp_path / "source.docx")
    plan = tmp_path / "operations.json"
    plan.write_text(
        json.dumps(
            {
                "version": "0.1",
                "operations": [
                    {
                        "type": "insert_section_before",
                        "target": _body_selector(),
                    },
                    {
                        "type": "set_footer_text",
                        "section_start": _body_selector(),
                        "text": "内部资料",
                        "alignment": "right",
                    },
                    {
                        "type": "set_paragraph_pagination",
                        "target": {
                            "text": "2 方案设计",
                            "role": "heading_1",
                            "after_role": "table_of_contents",
                        },
                        "page_break_before": True,
                        "keep_with_next": True,
                        "keep_together": True,
                        "widow_control": True,
                    },
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    first_output = tmp_path / "first.docx"
    second_output = tmp_path / "second.docx"

    first_report = apply_structure_operations_to_path(source, plan, first_output)
    second_report = apply_structure_operations_to_path(
        first_output, plan, second_output
    )
    structure = inspect_structure(first_output)
    document = Document(first_output)
    heading = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == "2 方案设计"
    )

    assert first_report["changed"] is True
    assert structure["sections"][0]["footer_text"] == ""
    assert structure["sections"][1]["footer_text"] == "内部资料"
    assert heading.paragraph_format.page_break_before is True
    assert heading.paragraph_format.keep_with_next is True
    assert heading.paragraph_format.keep_together is True
    assert heading.paragraph_format.widow_control is True
    assert second_report["changed"] is False


class _FakeFieldRefreshBackend:
    name = "fake-word"

    def __init__(self) -> None:
        self.called_with: Path | None = None

    def refresh(self, document_path: Path) -> dict[str, int]:
        self.called_with = document_path
        document = Document(document_path)
        document.core_properties.comments = "fields refreshed"
        document.save(document_path)
        return {
            "page_count": 3,
            "field_count": 4,
            "toc_count": 1,
        }


class _SectionCollapsingFieldRefreshBackend:
    name = "unsafe-word"

    def refresh(self, document_path: Path) -> dict[str, int]:
        document = Document(document_path)
        for section_properties in document.element.body.xpath(
            "./*[local-name()='p']/*[local-name()='pPr']"
            "/*[local-name()='sectPr']"
        ):
            section_properties.getparent().remove(section_properties)
        document.save(document_path)
        return {
            "page_count": 1,
            "field_count": 0,
            "toc_count": 0,
        }


def test_refresh_fields_uses_atomic_injected_backend(tmp_path: Path) -> None:
    source = _document_with_toc_and_body(tmp_path / "source.docx")
    output = tmp_path / "refreshed.docx"
    backend = _FakeFieldRefreshBackend()

    report = refresh_fields_to_path(source, output, backend=backend)

    assert output.exists()
    assert backend.called_with is not None
    assert backend.called_with != output
    assert report["backend"] == "fake-word"
    assert report["page_count"] == 3
    assert report["structure_preserved"] is True
    assert Document(output).core_properties.comments == "fields refreshed"


def test_refresh_fields_rejects_section_topology_changes(tmp_path: Path) -> None:
    source = _document_with_toc_and_body(tmp_path / "source.docx")
    plan = _operation_plan(tmp_path / "operations.json")
    structured = tmp_path / "structured.docx"
    output = tmp_path / "refreshed.docx"
    apply_structure_operations_to_path(source, plan, structured)

    with pytest.raises(FieldRefreshIntegrityError, match="结构指纹"):
        refresh_fields_to_path(
            structured,
            output,
            backend=_SectionCollapsingFieldRefreshBackend(),
        )

    assert not output.exists()


def test_public_api_exposes_structure_workflow(tmp_path: Path) -> None:
    source = _document_with_toc_and_body(tmp_path / "source.docx")
    plan = _operation_plan(tmp_path / "operations.json")
    expectations = _expectations(tmp_path / "expectations.json")
    output = tmp_path / "output.docx"

    structure = public_api.inspect_structure(source)
    operation_report = public_api.apply_structure_operations(
        source, plan, output
    )
    validation = public_api.validate_layout(output, expectations)

    assert structure["section_count"] == 1
    assert operation_report["changed"] is True
    assert validation["status"] == "passed"
    assert callable(public_api.refresh_fields)


def test_structure_cli_commands_write_reports(tmp_path: Path) -> None:
    source = _document_with_toc_and_body(tmp_path / "source.docx")
    plan = _operation_plan(tmp_path / "operations.json")
    expectations = _expectations(tmp_path / "expectations.json")
    structure_path = tmp_path / "structure.json"
    output = tmp_path / "output.docx"
    operation_report = tmp_path / "operations-report.json"
    validation_path = tmp_path / "validation.json"

    inspect_result = runner.invoke(
        app,
        ["inspect-structure", str(source), "-o", str(structure_path)],
    )
    apply_result = runner.invoke(
        app,
        [
            "apply-operations",
            str(source),
            "--plan",
            str(plan),
            "-o",
            str(output),
            "--report",
            str(operation_report),
        ],
    )
    validate_result = runner.invoke(
        app,
        [
            "validate-layout",
            str(output),
            "--expect",
            str(expectations),
            "-o",
            str(validation_path),
        ],
    )

    assert inspect_result.exit_code == 0
    assert apply_result.exit_code == 0
    assert validate_result.exit_code == 0
    assert json.loads(structure_path.read_text(encoding="utf-8"))[
        "section_count"
    ] == 1
    assert json.loads(operation_report.read_text(encoding="utf-8"))[
        "changed"
    ] is True
    assert json.loads(validation_path.read_text(encoding="utf-8"))[
        "status"
    ] == "passed"

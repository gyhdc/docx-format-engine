from __future__ import annotations

from pathlib import Path

from docx import Document

from word_format_tool.document_analyzer import (
    analyze_document,
    guess_paragraph_role,
)
from word_format_tool.rule_loader import load_rules


def test_analyzer_uses_patterns_before_style_names(
    draft_path: Path, rules_path: Path,
) -> None:
    analyzed = analyze_document(Document(draft_path), load_rules(rules_path))

    assert analyzed[0].role == "title_zh"
    assert analyzed[1].role == "heading_1"
    assert analyzed[2].role == "body"
    assert analyzed[3].role == "unknown"


def test_analyzer_recognizes_keywords_and_captions(
    tmp_path: Path, rules_path: Path,
) -> None:
    path = tmp_path / "roles.docx"
    document = Document()
    for text in ["摘要：内容", "关键词：岩土；智能", "图 1 结果", "表 1 参数", "[1] 文献"]:
        document.add_paragraph(text)
    document.save(path)

    roles = [
        item.role
        for item in analyze_document(Document(path), load_rules(rules_path))
    ]

    assert roles == [
        "abstract",
        "keywords_zh",
        "figure_caption",
        "table_caption",
        "reference_entry_zh",
    ]


def test_analyzer_skips_table_of_contents_entries(rules_path: Path) -> None:
    rules = load_rules(rules_path)

    entry_role = guess_paragraph_role("1 引言\t1", "toc 2", rules.detection)
    title_role = guess_paragraph_role("目    录", "Normal", rules.detection)

    assert entry_role == "unknown"
    assert title_role == "unknown"


def test_analyzer_tracks_reference_section_until_terminal_heading(
    tmp_path: Path,
    rules_path: Path,
) -> None:
    source = tmp_path / "reference-section.docx"
    document = Document()
    document.add_paragraph("1 引言")
    document.add_paragraph("正文内容。")
    document.add_paragraph("参考文献")
    document.add_paragraph("[1] 第一条文献")
    document.add_paragraph("未编号但仍属于文献区的条目")
    document.add_paragraph("致谢")
    document.add_paragraph("感谢老师。")
    document.save(source)

    analyzed = analyze_document(Document(source), load_rules(rules_path))

    assert [item.role for item in analyzed] == [
        "heading_1",
        "body",
        "reference_heading",
        "reference_entry_zh",
        "reference_entry_zh",
        "acknowledgements",
        "body",
    ]


def test_analyzer_recognizes_semantic_modules_when_styles_are_wrong(
    tmp_path: Path,
    rules_path: Path,
) -> None:
    source = tmp_path / "semantic-modules.docx"
    document = Document()
    for text in [
        "摘 要",
        "本文研究智能文档格式处理方法。",
        "关键词：格式识别；模板对齐",
        "ABSTRACT",
        "This paper studies semantic document formatting.",
        "Keywords: formatting; template",
        "目 录",
        "1 绪论........................1",
        "1.1 研究背景..................2",
        "1 绪论",
        "正文内容。",
        "参考文献",
        "[1] 第一条文献",
    ]:
        document.add_paragraph(text)
    document.save(source)

    roles = [
        item.role
        for item in analyze_document(Document(source), load_rules(rules_path))
    ]

    assert roles == [
        "abstract_heading_zh",
        "abstract_body_zh",
        "keywords_zh",
        "abstract_heading_en",
        "abstract_body_en",
        "keywords_en",
        "toc_heading",
        "toc_entry_1",
        "toc_entry_2",
        "heading_1",
        "body",
        "reference_heading",
        "reference_entry_zh",
    ]

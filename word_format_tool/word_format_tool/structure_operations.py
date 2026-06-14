"""Atomic, idempotent structural operations for Agent-authored plans."""

from __future__ import annotations

import json
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pydantic import ValidationError

from .document_io import (
    copy_docx_atomic,
    load_docx,
    prepare_output_path,
    save_docx_atomic,
)
from .exceptions import StructureOperationError
from .models import (
    InsertSectionBeforeOperation,
    RequestFieldUpdateOperation,
    SetFooterTextOperation,
    SetHeaderOperation,
    SetPageNumberOperation,
    SetParagraphPaginationOperation,
    StructureOperationPlan,
)
from .paragraph_locator import LocatedParagraph, locate_paragraph


ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "distribute": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
}
SECTION_BREAK_MARKER = "WORDFMT_SECTION_BREAK"


def _load_plan(path: str | Path) -> StructureOperationPlan:
    plan_path = Path(path)
    if not plan_path.is_file():
        raise StructureOperationError(f"操作计划不存在或不是文件: {plan_path}")
    try:
        payload = json.loads(plan_path.read_text(encoding="utf-8-sig"))
        return StructureOperationPlan.model_validate(payload)
    except (OSError, json.JSONDecodeError, ValidationError) as exc:
        raise StructureOperationError(f"操作计划无效 {plan_path}: {exc}") from exc


def _has_section_break_before(located: LocatedParagraph) -> bool:
    previous = located.paragraph._p.getprevious()
    return bool(
        previous is not None
        and previous.tag == qn("w:p")
        and previous.xpath("./*[local-name()='pPr']/*[local-name()='sectPr']")
    )


def _normalize_empty_section_break_carrier(
    located: LocatedParagraph,
) -> bool:
    """Make an empty break paragraph stable across Word TOC refreshes."""

    previous = located.paragraph._p.getprevious()
    if (
        previous is None
        or previous.tag != qn("w:p")
        or not previous.xpath(
            "./*[local-name()='pPr']/*[local-name()='sectPr']"
        )
    ):
        return False
    text = "".join(previous.xpath(".//*[local-name()='t']/text()"))
    non_marker_text = text.replace(SECTION_BREAK_MARKER, "").strip()
    non_text_content = previous.xpath(
        ".//*[local-name()='tab' or local-name()='br'"
        " or local-name()='drawing' or local-name()='object'"
        " or local-name()='pict' or local-name()='fldChar'"
        " or local-name()='instrText']"
    )
    if non_marker_text or non_text_content:
        return False
    properties = previous.find(qn("w:pPr"))
    if properties is None:
        return False
    changed = False
    for tag in (
        "w:pStyle",
        "w:outlineLvl",
        "w:keepNext",
        "w:keepLines",
        "w:pageBreakBefore",
    ):
        element = properties.find(qn(tag))
        if element is not None:
            properties.remove(element)
            changed = True
    marker_run = next(
        (
            run
            for run in previous.findall(qn("w:r"))
            if SECTION_BREAK_MARKER
            in "".join(
                text_element.text or ""
                for text_element in run.findall(qn("w:t"))
            )
        ),
        None,
    )
    if marker_run is None:
        marker_run = OxmlElement("w:r")
        marker_text = OxmlElement("w:t")
        marker_text.text = SECTION_BREAK_MARKER
        marker_run.append(marker_text)
        previous.append(marker_run)
        changed = True
    run_properties = marker_run.find(qn("w:rPr"))
    if run_properties is None:
        run_properties = OxmlElement("w:rPr")
        marker_run.insert(0, run_properties)
    if run_properties.find(qn("w:vanish")) is None:
        run_properties.append(OxmlElement("w:vanish"))
        changed = True
    return changed


def _find_outer_toc_end_in_target(
    document: Any, located: LocatedParagraph,
) -> Any | None:
    field_stack: list[dict[str, str]] = []
    for child in document.element.body.iterchildren():
        if child.tag != qn("w:p"):
            continue
        for element in child.iter():
            if element.tag == qn("w:fldChar"):
                field_type = element.get(qn("w:fldCharType"))
                if field_type == "begin":
                    field_stack.append({"instruction": ""})
                elif field_type == "end" and field_stack:
                    field = field_stack.pop()
                    if (
                        child is located.paragraph._p
                        and field["instruction"].strip().upper().startswith(
                            "TOC "
                        )
                    ):
                        return element
            elif element.tag == qn("w:instrText") and field_stack:
                field_stack[-1]["instruction"] += element.text or ""
        if child is located.paragraph._p:
            break
    return None


def _move_outer_toc_end_before_section(
    document: Any, located: LocatedParagraph,
) -> bool:
    field_end = _find_outer_toc_end_in_target(document, located)
    if field_end is None:
        return False
    carrier = located.paragraph._p.getprevious()
    destination = carrier.getprevious() if carrier is not None else None
    if destination is None or destination.tag != qn("w:p"):
        destination = OxmlElement("w:p")
        carrier.addprevious(destination)
    destination_run = OxmlElement("w:r")
    destination_run.append(deepcopy(field_end))
    destination.append(destination_run)
    source_run = field_end.getparent()
    source_run.remove(field_end)
    if (
        source_run.tag == qn("w:r")
        and not any(child.tag != qn("w:rPr") for child in source_run)
    ):
        source_run.getparent().remove(source_run)
    return True


def _set_section_type(sect_pr: Any, value: str) -> None:
    types = sect_pr.xpath("./*[local-name()='type']")
    if types:
        element = types[0]
    else:
        element = OxmlElement("w:type")
        sect_pr.insert(0, element)
    element.set(qn("w:val"), value)


def _insert_section_before(
    document: Any, operation: InsertSectionBeforeOperation,
) -> tuple[bool, dict[str, Any]]:
    located = locate_paragraph(document, operation.target)
    if _has_section_break_before(located):
        moved_toc_end = _move_outer_toc_end_before_section(document, located)
        normalized = _normalize_empty_section_break_carrier(located)
        return moved_toc_end or normalized, {
            "target": located.location,
            "section_count": len(document.sections),
            "moved_toc_field_end": moved_toc_end,
            "normalized_empty_carrier": normalized,
            "postcondition_passed": True,
        }
    trailing_sect_pr = document.sections[-1]._sectPr
    sect_pr = deepcopy(trailing_sect_pr)
    _set_section_type(sect_pr, "nextPage")
    for reference in trailing_sect_pr.xpath(
        "./*[local-name()='headerReference' or local-name()='footerReference']"
    ):
        trailing_sect_pr.remove(reference)
    break_paragraph = OxmlElement("w:p")
    properties = OxmlElement("w:pPr")
    properties.append(sect_pr)
    break_paragraph.append(properties)
    located.paragraph._p.addprevious(break_paragraph)
    refreshed = locate_paragraph(document, operation.target)
    moved_toc_end = _move_outer_toc_end_before_section(document, refreshed)
    _normalize_empty_section_break_carrier(refreshed)
    passed = _has_section_break_before(refreshed)
    if not passed:
        raise StructureOperationError(
            f"在 {refreshed.location} 前插入分节符后验证失败。"
        )
    return True, {
        "target": refreshed.location,
        "section_count": len(document.sections),
        "moved_toc_field_end": moved_toc_end,
        "postcondition_passed": passed,
    }


def _section_index_for(document: Any, located: LocatedParagraph) -> int:
    section_index = 0
    for child in document.element.body.iterchildren():
        if child is located.paragraph._p:
            return section_index
        if (
            child.tag == qn("w:p")
            and child.xpath("./*[local-name()='pPr']/*[local-name()='sectPr']")
        ):
            section_index += 1
    raise StructureOperationError(f"无法确定 {located.location} 所属分节。")


def _clear_paragraph(paragraph: Any) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _set_bottom_border(paragraph: Any, enabled: bool) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if not enabled:
        if borders is not None:
            properties.remove(borders)
        return
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")


def _has_bottom_border(paragraph: Any) -> bool:
    values = paragraph._p.xpath(
        "./*[local-name()='pPr']/*[local-name()='pBdr']"
        "/*[local-name()='bottom']/@*[local-name()='val']"
    )
    return bool(values and values[0] not in {"nil", "none"})


def _set_header(
    document: Any, operation: SetHeaderOperation,
) -> tuple[bool, dict[str, Any]]:
    located = locate_paragraph(document, operation.section_start)
    section_index = _section_index_for(document, located)
    section = document.sections[section_index]
    header = section.header
    was_linked = header.is_linked_to_previous
    if was_linked:
        header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    desired_alignment = ALIGNMENTS[operation.alignment]
    current_text = "\n".join(
        item.text.strip() for item in header.paragraphs if item.text.strip()
    )
    already = (
        not was_linked
        and current_text == operation.text
        and paragraph.alignment == desired_alignment
        and _has_bottom_border(paragraph) == operation.bottom_border
    )
    if not already:
        paragraph = header.paragraphs[0]
        _clear_paragraph(paragraph)
        paragraph.add_run(operation.text)
        paragraph.alignment = desired_alignment
        _set_bottom_border(paragraph, operation.bottom_border)
    passed = (
        not header.is_linked_to_previous
        and header.paragraphs[0].text == operation.text
        and header.paragraphs[0].alignment == desired_alignment
    )
    if not passed:
        raise StructureOperationError(
            f"设置 section[{section_index}] 页眉后验证失败。"
        )
    return not already, {
        "section_index": section_index,
        "text": header.paragraphs[0].text,
        "postcondition_passed": passed,
    }


def _field_codes(element: Any) -> list[str]:
    codes = [
        str(text).strip()
        for text in element.xpath(".//*[local-name()='instrText']/text()")
        if str(text).strip()
    ]
    codes.extend(
        str(value).strip()
        for value in element.xpath(
            ".//*[local-name()='fldSimple']/@*[local-name()='instr']"
        )
        if str(value).strip()
    )
    return codes


def _page_start(section: Any) -> int | None:
    values = section._sectPr.xpath(
        "./*[local-name()='pgNumType']/@*[local-name()='start']"
    )
    return int(values[0]) if values else None


def _set_page_start(section: Any, start: int) -> None:
    values = section._sectPr.xpath("./*[local-name()='pgNumType']")
    if values:
        page_number = values[0]
    else:
        page_number = OxmlElement("w:pgNumType")
        section._sectPr.append(page_number)
    page_number.set(qn("w:start"), str(start))


def _set_page_number(
    document: Any, operation: SetPageNumberOperation,
) -> tuple[bool, dict[str, Any]]:
    located = locate_paragraph(document, operation.section_start)
    section_index = _section_index_for(document, located)
    section = document.sections[section_index]
    footer = section.footer
    was_linked = footer.is_linked_to_previous
    if was_linked:
        footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    desired_alignment = ALIGNMENTS[operation.alignment]
    page_codes = [
        code for code in _field_codes(footer._element) if "PAGE" in code.upper()
    ]
    already = (
        not was_linked
        and len(page_codes) == 1
        and _page_start(section) == operation.start
        and paragraph.alignment == desired_alignment
    )
    if not already:
        paragraph = footer.paragraphs[0]
        _clear_paragraph(paragraph)
        paragraph.alignment = desired_alignment
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), " PAGE ")
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = str(operation.start)
        run.append(text)
        field.append(run)
        paragraph._p.append(field)
        _set_page_start(section, operation.start)
    page_codes = [
        code for code in _field_codes(footer._element) if "PAGE" in code.upper()
    ]
    passed = (
        not footer.is_linked_to_previous
        and len(page_codes) == 1
        and _page_start(section) == operation.start
    )
    if not passed:
        raise StructureOperationError(
            f"设置 section[{section_index}] 页码后验证失败。"
        )
    return not already, {
        "section_index": section_index,
        "page_number_start": _page_start(section),
        "page_field_count": len(page_codes),
        "postcondition_passed": passed,
    }


def _set_footer_text(
    document: Any, operation: SetFooterTextOperation,
) -> tuple[bool, dict[str, Any]]:
    located = locate_paragraph(document, operation.section_start)
    section_index = _section_index_for(document, located)
    section = document.sections[section_index]
    footer = section.footer
    was_linked = footer.is_linked_to_previous
    if was_linked:
        footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    desired_alignment = ALIGNMENTS[operation.alignment]
    current_text = "\n".join(
        item.text.strip() for item in footer.paragraphs if item.text.strip()
    )
    already = (
        not was_linked
        and current_text == operation.text
        and paragraph.alignment == desired_alignment
        and not _field_codes(footer._element)
    )
    if not already:
        _clear_paragraph(paragraph)
        paragraph.add_run(operation.text)
        paragraph.alignment = desired_alignment
    passed = (
        not footer.is_linked_to_previous
        and footer.paragraphs[0].text == operation.text
        and footer.paragraphs[0].alignment == desired_alignment
    )
    if not passed:
        raise StructureOperationError(
            f"设置 section[{section_index}] 页脚文本后验证失败。"
        )
    return not already, {
        "section_index": section_index,
        "text": footer.paragraphs[0].text,
        "postcondition_passed": passed,
    }


def _set_paragraph_pagination(
    document: Any, operation: SetParagraphPaginationOperation,
) -> tuple[bool, dict[str, Any]]:
    located = locate_paragraph(document, operation.target)
    paragraph_format = located.paragraph.paragraph_format
    desired = {
        "page_break_before": operation.page_break_before,
        "keep_with_next": operation.keep_with_next,
        "keep_together": operation.keep_together,
        "widow_control": operation.widow_control,
    }
    selected = {key: value for key, value in desired.items() if value is not None}
    before = {key: getattr(paragraph_format, key) for key in selected}
    changed = any(before[key] is not value for key, value in selected.items())
    for key, value in selected.items():
        setattr(paragraph_format, key, value)
    after = {key: getattr(paragraph_format, key) for key in selected}
    passed = all(after[key] is value for key, value in selected.items())
    if not passed:
        raise StructureOperationError(
            f"设置 {located.location} 分页属性后验证失败。"
        )
    return changed, {
        "target": located.location,
        "before": before,
        "after": after,
        "postcondition_passed": passed,
    }


def _request_field_update(
    document: Any, operation: RequestFieldUpdateOperation,
) -> tuple[bool, dict[str, Any]]:
    del operation
    settings = document.settings.element
    values = settings.xpath("./*[local-name()='updateFields']")
    if values:
        element = values[0]
        already = element.get(qn("w:val")) in {"true", "1", "on"}
    else:
        element = OxmlElement("w:updateFields")
        settings.append(element)
        already = False
    element.set(qn("w:val"), "true")
    return not already, {
        "update_fields_enabled": True,
        "postcondition_passed": True,
    }


def apply_structure_operations_to_path(
    document_path: str | Path,
    plan_path: str | Path,
    output_path: str | Path,
) -> dict[str, Any]:
    """Apply a strict operation plan and publish only after all checks pass."""

    document, source = load_docx(document_path)
    output = prepare_output_path(
        output_path, protected_input=source, suffix=".docx"
    )
    plan = _load_plan(plan_path)
    results: list[dict[str, Any]] = []
    changed = False
    for operation in plan.operations:
        if isinstance(operation, InsertSectionBeforeOperation):
            operation_changed, details = _insert_section_before(document, operation)
        elif isinstance(operation, SetHeaderOperation):
            operation_changed, details = _set_header(document, operation)
        elif isinstance(operation, SetPageNumberOperation):
            operation_changed, details = _set_page_number(document, operation)
        elif isinstance(operation, SetFooterTextOperation):
            operation_changed, details = _set_footer_text(document, operation)
        elif isinstance(operation, SetParagraphPaginationOperation):
            operation_changed, details = _set_paragraph_pagination(
                document, operation
            )
        elif isinstance(operation, RequestFieldUpdateOperation):
            operation_changed, details = _request_field_update(document, operation)
        else:  # pragma: no cover - discriminated schema prevents this branch
            raise StructureOperationError(
                f"不支持的结构操作: {operation.type}"
            )
        changed = changed or operation_changed
        results.append(
            {
                "type": operation.type,
                "changed": operation_changed,
                **details,
            }
        )
    if changed:
        save_docx_atomic(document, output)
    else:
        copy_docx_atomic(source, output)
    return {
        "source": str(source),
        "output": str(output),
        "plan": str(plan_path),
        "changed": changed,
        "operations": results,
    }

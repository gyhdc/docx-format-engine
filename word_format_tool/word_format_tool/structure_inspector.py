"""Machine-readable DOCX structure inspection."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx.oxml.ns import qn

from .document_analyzer import analyze_document
from .document_io import load_docx
from .models import AnalyzedParagraph, Rules
from .ooxml_utils import section_to_dict
from .paragraph_locator import normalize_paragraph_text


TOC_ROLES = {
    "table_of_contents",
    "toc_heading",
    "toc_entry_1",
    "toc_entry_2",
    "toc_entry_3",
}


def _analysis_rules() -> Rules:
    return Rules.model_validate({"version": "0.1", "roles": {"body": {}}})


def _field_codes(element: Any) -> list[str]:
    codes = [
        normalize_paragraph_text(text)
        for text in element.xpath(".//*[local-name()='instrText']/text()")
        if normalize_paragraph_text(text)
    ]
    for field in element.xpath(".//*[local-name()='fldSimple']"):
        instruction = field.get(qn("w:instr"))
        if instruction and normalize_paragraph_text(instruction):
            codes.append(normalize_paragraph_text(instruction))
    return codes


def _page_number_start(section: Any) -> int | None:
    values = section._sectPr.xpath(
        "./*[local-name()='pgNumType']/@*[local-name()='start']"
    )
    if not values:
        return None
    try:
        return int(values[0])
    except (TypeError, ValueError):
        return None


def _section_break_type(section: Any) -> str:
    values = section._sectPr.xpath(
        "./*[local-name()='type']/@*[local-name()='val']"
    )
    return str(values[0]) if values else "nextPage"


def _paragraph_section_indexes(document: Any) -> list[int]:
    indexes: list[int] = []
    section_index = 0
    for child in document.element.body.iterchildren():
        if child.tag != qn("w:p"):
            continue
        indexes.append(section_index)
        if child.xpath("./*[local-name()='pPr']/*[local-name()='sectPr']"):
            section_index += 1
    if len(indexes) != len(document.paragraphs):
        return [0] * len(document.paragraphs)
    return indexes


def _story_text(part: Any) -> str:
    return "\n".join(
        paragraph.text.strip()
        for paragraph in part.paragraphs
        if paragraph.text.strip()
    )


def _toc_items(analyzed: list[AnalyzedParagraph]) -> list[AnalyzedParagraph]:
    return [
        item
        for item in analyzed
        if item.story == "body"
        and (
            item.ignore_reason == "table_of_contents"
            or item.role in TOC_ROLES
        )
    ]


def inspect_structure(document_path: str | Path) -> dict[str, Any]:
    """Inspect structural facts without invoking Word or modifying the DOCX."""

    document, source = load_docx(document_path)
    analyzed = analyze_document(document, _analysis_rules())
    body_items = [item for item in analyzed if item.story == "body"]
    section_indexes = _paragraph_section_indexes(document)
    toc_items = _toc_items(analyzed)
    toc_start = min((item.paragraph_index for item in toc_items), default=None)
    toc_end = max((item.paragraph_index for item in toc_items), default=None)

    meaningful_by_section: dict[int, list[int]] = {}
    for item in body_items:
        if not item.text.strip():
            continue
        section_index = section_indexes[item.paragraph_index]
        meaningful_by_section.setdefault(section_index, []).append(
            item.paragraph_index
        )

    paragraphs: list[dict[str, Any]] = []
    for item in body_items:
        section_index = section_indexes[item.paragraph_index]
        first_in_section = (
            meaningful_by_section.get(section_index, [None])[0]
            == item.paragraph_index
        )
        after_toc = toc_end is not None and item.paragraph_index > toc_end
        locator: dict[str, Any] = {
            "text": item.text.strip(),
            "role": item.role,
        }
        if after_toc:
            locator["after_role"] = "table_of_contents"
        paragraphs.append(
            {
                "paragraph_index": item.paragraph_index,
                "text": item.text,
                "role": item.role,
                "style_name": item.style_name,
                "location": item.location,
                "section_index": section_index,
                "starts_section": bool(item.text.strip() and first_in_section),
                "page_break_before": bool(
                    item.format.get("page_break_before", False)
                ),
                "keep_with_next": bool(item.format.get("keep_with_next", False)),
                "ignore_reason": item.ignore_reason,
                "locator": locator,
            }
        )

    sections: list[dict[str, Any]] = []
    for section_index, section in enumerate(document.sections):
        indexes = meaningful_by_section.get(section_index, [])
        sections.append(
            {
                "index": section_index,
                "break_type": _section_break_type(section),
                "page": section_to_dict(section),
                "start_paragraph": indexes[0] if indexes else None,
                "end_paragraph": indexes[-1] if indexes else None,
                "header_text": _story_text(section.header),
                "footer_text": _story_text(section.footer),
                "header_linked_to_previous": section.header.is_linked_to_previous,
                "footer_linked_to_previous": section.footer.is_linked_to_previous,
                "header_field_codes": _field_codes(section.header._element),
                "footer_field_codes": _field_codes(section.footer._element),
                "page_number_start": _page_number_start(section),
            }
        )

    heading_after_toc = next(
        (
            item
            for item in body_items
            if item.role == "heading_1"
            and (toc_end is None or item.paragraph_index > toc_end)
        ),
        None,
    )
    update_values = document.settings.element.xpath(
        "./*[local-name()='updateFields']/@*[local-name()='val']"
    )
    return {
        "document": str(source),
        "section_count": len(document.sections),
        "sections": sections,
        "paragraphs": paragraphs,
        "toc": {
            "present": bool(toc_items),
            "start_paragraph": toc_start,
            "end_paragraph": toc_end,
            "field_codes": _field_codes(document.element.body),
        },
        "body_first_heading": (
            {
                "paragraph_index": heading_after_toc.paragraph_index,
                "text": heading_after_toc.text,
                "role": heading_after_toc.role,
            }
            if heading_after_toc is not None
            else None
        ),
        "update_fields_enabled": any(
            str(value).casefold() in {"true", "1", "on"} for value in update_values
        ),
        "word_pagination_available": False,
    }

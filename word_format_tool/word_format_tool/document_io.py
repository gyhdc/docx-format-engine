"""Small, shared boundary for DOCX path validation, loading, and saving.

Purpose:
    Keep filesystem and package-read errors consistent across public APIs.
MVP scope:
    Handles local ``.docx`` files only and creates missing output directories.
"""

from __future__ import annotations

import hashlib
import os
import shutil
import tempfile
import zipfile
from collections.abc import Callable
from pathlib import Path
from typing import Any

from docx import Document

from .exceptions import (
    DocumentReadError,
    InputFileError,
    UnsafeOutputPathError,
    WordFormatToolError,
)


def require_docx_file(path: str | Path) -> Path:
    """Validate and return one existing DOCX path."""

    resolved = Path(path)
    if not resolved.is_file():
        raise InputFileError(f"DOCX 文件不存在或不是文件: {resolved}")
    if resolved.suffix.lower() != ".docx":
        raise InputFileError(f"仅支持 .docx 文件: {resolved}")
    return resolved


def load_docx(path: str | Path) -> tuple[Any, Path]:
    """Open a DOCX and preserve the original exception as the cause."""

    docx_path = require_docx_file(path)
    try:
        return Document(docx_path), docx_path
    except Exception as exc:
        raise DocumentReadError(f"无法读取 DOCX {docx_path}: {exc}") from exc


def prepare_output_path(
    output_path: str | Path,
    *,
    protected_input: str | Path | None = None,
    suffix: str | None = None,
) -> Path:
    """Validate an output path and create its parent directory."""

    output = Path(output_path)
    if suffix is not None and output.suffix.lower() != suffix.lower():
        raise InputFileError(f"输出文件必须使用 {suffix} 扩展名: {output}")
    if protected_input is not None:
        source = Path(protected_input)
        if output.resolve() == source.resolve():
            raise UnsafeOutputPathError(
                f"拒绝覆盖原始输入文件，请指定不同输出路径: {output}"
            )
    try:
        output.parent.mkdir(parents=True, exist_ok=True)
    except OSError as exc:
        raise InputFileError(f"无法创建输出目录 {output.parent}: {exc}") from exc
    return output


def save_docx(document: Any, output_path: str | Path) -> Path:
    """Save one DOCX with a clear domain error."""

    output = Path(output_path)
    try:
        document.save(output)
    except Exception as exc:
        raise DocumentReadError(f"无法写入 DOCX {output}: {exc}") from exc
    return output


def _content_fingerprint(path: str | Path) -> dict[str, Any]:
    document, docx_path = load_docx(path)
    body = document.element.body
    body_text = "".join(
        node.text or "" for node in body.xpath(".//*[local-name()='t']")
    )
    body_block_order = tuple(
        child.tag.rsplit("}", 1)[-1] for child in body.iterchildren()
    )
    tables = tuple(
        (
            len(table.rows),
            len(table.columns),
            tuple(
                tuple(cell.text for cell in row.cells)
                for row in table.rows
            ),
        )
        for table in document.tables
    )
    formulas = tuple(
        "".join(element.itertext())
        for element in document.element.xpath(".//*[local-name()='oMath']")
    )
    try:
        with zipfile.ZipFile(docx_path) as package:
            media = tuple(
                (
                    name,
                    hashlib.sha256(package.read(name)).hexdigest(),
                )
                for name in sorted(package.namelist())
                if name.startswith("word/media/") and not name.endswith("/")
            )
    except (OSError, zipfile.BadZipFile, KeyError) as exc:
        raise DocumentReadError(
            f"无法校验 DOCX 媒体内容 {docx_path}: {exc}"
        ) from exc

    layout_fields = {
        "anchor",
        "inline",
        "positionH",
        "positionV",
        "posOffset",
        "align",
        "extent",
        "effectExtent",
        "simplePos",
        "wrapNone",
        "wrapSquare",
        "wrapTight",
        "wrapThrough",
        "wrapTopAndBottom",
        "docPr",
        "blip",
        "shape",
        "pict",
        "txbxContent",
    }
    story_roots = [document.element]
    seen_parts: set[Any] = set()
    for section in document.sections:
        for part in (section.header, section.footer):
            if part._element in seen_parts:
                continue
            seen_parts.add(part._element)
            story_roots.append(part._element)
    layout_objects: list[tuple[Any, ...]] = []
    for story_root in story_roots:
        roots = story_root.xpath(
            ".//*[local-name()='anchor' or local-name()='inline' "
            "or local-name()='pict' or "
            "(local-name()='shape' and namespace-uri()="
            "'urn:schemas-microsoft-com:vml')]"
        )
        for root in roots:
            nodes: list[tuple[Any, ...]] = []
            for node in root.iter():
                local_name = node.tag.rsplit("}", 1)[-1]
                if local_name not in layout_fields:
                    continue
                attributes = tuple(
                    sorted(
                        (
                            key.rsplit("}", 1)[-1],
                            value,
                        )
                        for key, value in node.attrib.items()
                    )
                )
                text = (node.text or "").strip()
                nodes.append((local_name, attributes, text))
            layout_objects.append(tuple(nodes))
    return {
        "body_text": body_text,
        "body_block_order": body_block_order,
        "paragraph_text": tuple(
            paragraph.text for paragraph in document.paragraphs
        ),
        "tables": tables,
        "formulas": formulas,
        "media": media,
        "layout_objects": tuple(layout_objects),
        "header_footer_text": tuple(
            (
                section_index,
                story,
                tuple(paragraph.text for paragraph in part.paragraphs),
            )
            for section_index, section in enumerate(document.sections)
            for story, part in (
                ("header", section.header),
                ("footer", section.footer),
            )
        ),
        "bookmark_names": tuple(
            sorted(
                name
                for name in document.element.xpath(
                    ".//*[local-name()='bookmarkStart']"
                    "/@*[local-name()='name']"
                )
                if name
            )
        ),
        "internal_hyperlink_anchors": tuple(
            sorted(
                anchor
                for anchor in document.element.xpath(
                    ".//*[local-name()='hyperlink']"
                    "/@*[local-name()='anchor']"
                )
                if anchor
            )
        ),
    }


def validate_preserved_content(
    source_path: str | Path,
    candidate_path: str | Path,
) -> None:
    """Raise when a repaired DOCX changes protected document content."""

    source = _content_fingerprint(source_path)
    candidate = _content_fingerprint(candidate_path)
    navigation_fields = {"bookmark_names", "internal_hyperlink_anchors"}
    changed = [
        field for field, value in source.items()
        if field not in navigation_fields and candidate.get(field) != value
    ]
    if changed:
        raise DocumentReadError(
            "DOCX 内容完整性校验失败，以下内容发生变化: "
            + ", ".join(changed)
        )
    allowed_prefixes = ("_SDAU_REF_", "_SDAU_CITE_", "_SDAU_TOC_")
    for field in sorted(navigation_fields):
        source_values = set(source[field])
        candidate_values = set(candidate[field])
        missing = sorted(source_values - candidate_values)
        unexpected = sorted(
            value
            for value in candidate_values - source_values
            if not value.startswith(allowed_prefixes)
        )
        if missing or unexpected:
            details = []
            if missing:
                details.append(f"missing={missing}")
            if unexpected:
                details.append(f"unexpected={unexpected}")
            raise DocumentReadError(
                "DOCX navigation 完整性校验失败，"
                f"{field}: {'; '.join(details)}"
            )


def _temporary_docx_path(output: Path) -> Path:
    file_descriptor, temporary_name = tempfile.mkstemp(
        prefix=f".{output.stem}-",
        suffix=".tmp.docx",
        dir=output.parent,
    )
    os.close(file_descriptor)
    return Path(temporary_name)


def _publish_docx_atomically(
    output: Path,
    writer: Callable[[Path], None],
    validator: Callable[[Path], None] | None = None,
) -> Path:
    temporary = _temporary_docx_path(output)
    try:
        writer(temporary)
        load_docx(temporary)
        if validator is not None:
            validator(temporary)
        os.replace(temporary, output)
    except WordFormatToolError:
        raise
    except Exception as exc:
        raise DocumentReadError(
            f"无法安全写入 DOCX {output}: {exc}"
        ) from exc
    finally:
        try:
            temporary.unlink(missing_ok=True)
        except OSError:
            pass
    return output


def save_docx_atomic(
    document: Any,
    output_path: str | Path,
    *,
    validator: Callable[[Path], None] | None = None,
) -> Path:
    """Serialize, validate, and atomically publish one DOCX."""

    output = Path(output_path)
    return _publish_docx_atomically(
        output,
        lambda temporary: document.save(temporary),
        validator,
    )


def copy_docx_atomic(
    source_path: str | Path,
    output_path: str | Path,
) -> Path:
    """Atomically copy a valid DOCX without reserializing its package."""

    source = require_docx_file(source_path)
    output = Path(output_path)
    return _publish_docx_atomically(
        output,
        lambda temporary: shutil.copyfile(source, temporary),
    )

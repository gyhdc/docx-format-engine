"""Compare analyzed Word formatting with validated rules and build issues.

Purpose:
    Perform deterministic paragraph, page, and basic caption-position checks.
MVP scope:
    Formatting and margins are fixable. Figure/table caption positions are
    inspected from OOXML block adjacency but never moved automatically.
"""

from __future__ import annotations

import re
from collections import Counter
from collections.abc import Callable
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

from .constants import (
    FONT_SIZE_TOLERANCE_PT,
    INDENT_TOLERANCE_CHARS,
    MARGIN_TOLERANCE_CM,
    SPACING_TOLERANCE_PT,
    cm_to_chars,
)
from .layout_safety import LayoutRiskProfile, analyze_layout_risk
from .models import (
    AnalyzedParagraph,
    CoverageSummary,
    FormatIssue,
    FormatReport,
    ReportSummary,
    RoleFormatRule,
    Rules,
)
from .ooxml_utils import section_to_dict
from .reference_links import (
    analyze_reference_document,
    inspect_reference_navigation,
)
from .rule_loader import resolve_role_rule
from .toc_repair import find_toc_heading_match

TOC_FIELD_ERROR_PATTERN = re.compile(
    r"(?:错误|error)[!！]\s*(?:"
    r"未定义书签|未找到引用源|"
    r"bookmark not defined|reference source not found"
    r")",
    flags=re.IGNORECASE,
)
PAGEREF_PATTERN = re.compile(
    r"\bPAGEREF\s+(?:\"([^\"]+)\"|([^\s\\]+))",
    flags=re.IGNORECASE,
)


@dataclass
class _IssueFactory:
    issues: list[FormatIssue]

    def add(
        self,
        *,
        severity: str,
        paragraph_index: int | None,
        role: str,
        text_preview: str,
        field: str,
        expected: Any,
        actual: Any,
        message: str,
        fixable: bool,
        approximation: str | None = None,
        location: str | None = None,
    ) -> None:
        self.issues.append(
            FormatIssue(
                id=f"ISSUE-{len(self.issues) + 1:04d}",
                severity=severity,
                paragraph_index=paragraph_index,
                role=role,
                text_preview=text_preview[:80],
                field=field,
                expected=expected,
                actual=actual,
                message=message,
                fixable=fixable,
                approximation=approximation,
                location=location,
            )
        )


def _same_string(expected: str, actual: Any) -> bool:
    return isinstance(actual, str) and expected.casefold() == actual.casefold()


def _close(expected: float, actual: Any, tolerance: float) -> bool:
    return isinstance(actual, (int, float)) and abs(expected - float(actual)) <= tolerance


def _display(value: Any) -> str:
    if value is None:
        return "未设置或无法解析"
    if isinstance(value, float):
        return f"{value:.4g}"
    return str(value)


def _check_simple_field(
    factory: _IssueFactory,
    item: AnalyzedParagraph,
    field: str,
    expected: Any,
    comparator: Callable[[Any, Any], bool],
) -> None:
    if expected is None:
        return
    actual = item.format.get(field)
    if comparator(expected, actual):
        return
    factory.add(
        severity="error",
        paragraph_index=item.paragraph_index,
        role=item.role,
        text_preview=item.text,
        field=field,
        expected=expected,
        actual=actual,
        message=(
            f"{item.role} 段落的 {field} 应为 {_display(expected)}，"
            f"实际为 {_display(actual)}。"
        ),
        fixable=True,
        location=item.location,
    )


def _check_indent(
    factory: _IssueFactory,
    item: AnalyzedParagraph,
    rule: RoleFormatRule,
    rule_field: str,
    actual_cm_field: str,
) -> None:
    expected = getattr(rule, rule_field)
    if expected is None:
        return
    actual_cm = float(item.format.get(actual_cm_field) or 0.0)
    actual_chars = round(
        cm_to_chars(actual_cm, item.format.get("font_size_pt")), 4
    )
    if abs(expected - actual_chars) <= INDENT_TOLERANCE_CHARS:
        return
    factory.add(
        severity="error",
        paragraph_index=item.paragraph_index,
        role=item.role,
        text_preview=item.text,
        field=rule_field,
        expected=expected,
        actual=actual_chars,
        message=(
            f"{item.role} 段落的 {rule_field} 应约为 {expected} 字符，"
            f"实际约为 {actual_chars} 字符。"
        ),
        fixable=True,
        approximation="字符缩进按代表字号换算为物理长度，属于近似值。",
        location=item.location,
    )


def _check_paragraph_rule(
    factory: _IssueFactory,
    item: AnalyzedParagraph,
    rule: RoleFormatRule,
) -> None:
    _check_simple_field(factory, item, "font_east_asia", rule.font_east_asia, _same_string)
    _check_simple_field(factory, item, "font_ascii", rule.font_ascii, _same_string)
    _check_simple_field(
        factory,
        item,
        "font_size_pt",
        rule.font_size_pt,
        lambda expected, actual: _close(expected, actual, FONT_SIZE_TOLERANCE_PT),
    )
    _check_simple_field(factory, item, "bold", rule.bold, lambda a, b: a is b)
    _check_simple_field(factory, item, "italic", rule.italic, lambda a, b: a is b)
    _check_simple_field(
        factory, item, "alignment", rule.alignment, lambda a, b: a == b
    )
    if rule.line_spacing is not None:
        actual = item.format.get("line_spacing")
        correct = (
            item.format.get("line_spacing_kind") == "multiple"
            and _close(rule.line_spacing, actual, 0.02)
        )
        if not correct:
            factory.add(
                severity="error",
                paragraph_index=item.paragraph_index,
                role=item.role,
                text_preview=item.text,
                field="line_spacing",
                expected=rule.line_spacing,
                actual=actual,
                message=(
                    f"{item.role} 段落行距应为 {rule.line_spacing} 倍，"
                    f"实际为 {_display(actual)}"
                    f"（{item.format.get('line_spacing_kind')}）。"
                ),
                fixable=True,
                location=item.location,
            )
    _check_simple_field(
        factory,
        item,
        "left_indent_cm",
        rule.left_indent_cm,
        lambda expected, actual: _close(expected, actual, MARGIN_TOLERANCE_CM),
    )
    _check_indent(
        factory,
        item,
        rule,
        "first_line_indent_chars",
        "first_line_indent_cm",
    )
    _check_indent(
        factory,
        item,
        rule,
        "hanging_indent_chars",
        "hanging_indent_cm",
    )
    _check_simple_field(
        factory,
        item,
        "space_before_pt",
        rule.space_before_pt,
        lambda expected, actual: _close(expected, actual, SPACING_TOLERANCE_PT),
    )
    _check_simple_field(
        factory,
        item,
        "space_after_pt",
        rule.space_after_pt,
        lambda expected, actual: _close(expected, actual, SPACING_TOLERANCE_PT),
    )
    _check_simple_field(
        factory,
        item,
        "keep_with_next",
        rule.keep_with_next,
        lambda a, b: a is b,
    )


def _check_page_rules(
    factory: _IssueFactory,
    document: Any,
    rules: Rules,
    layout_risk: LayoutRiskProfile,
) -> None:
    if rules.page is None:
        return
    expected = rules.page.model_dump(exclude_none=True)
    paper_sizes = {
        "A4": (21.0, 29.7),
        "Letter": (21.59, 27.94),
        "Legal": (21.59, 35.56),
    }
    if "paper_size" in expected:
        width, height = paper_sizes[expected.pop("paper_size")]
        expected.setdefault("page_width_cm", width)
        expected.setdefault("page_height_cm", height)

    geometry_fixable = (
        not layout_risk.high_risk or rules.priority.allow_unsafe_page_geometry
    )
    mismatch_found = False
    for section_index, section in enumerate(document.sections):
        actual = section_to_dict(section)
        for field, expected_value in expected.items():
            actual_value = actual.get(field)
            if _close(expected_value, actual_value, MARGIN_TOLERANCE_CM):
                continue
            mismatch_found = True
            factory.add(
                severity="error",
                paragraph_index=None,
                role="document",
                text_preview=f"Section {section_index + 1}",
                field=field,
                expected=expected_value,
                actual=actual_value,
                message=(
                    f"第 {section_index + 1} 节的 {field} 应为 "
                    f"{expected_value} cm，实际为 {_display(actual_value)} cm。"
                ),
                fixable=geometry_fixable,
            )
    if mismatch_found and layout_risk.high_risk:
        factory.add(
            severity="warning",
            paragraph_index=None,
            role="document",
            text_preview="页面布局风险",
            field="page_layout_risk",
            expected="修改页面几何前无封面表格、浮动对象或文本框风险",
            actual=layout_risk.to_dict(),
            message=(
                "检测到可能依赖当前页面几何的对象；页面尺寸和边距默认只报告。"
                + (
                    "调用方已显式允许高风险页面几何修改。"
                    if rules.priority.allow_unsafe_page_geometry
                    else "如确需修改，必须显式设置 allow_unsafe_page_geometry。"
                )
            ),
            fixable=False,
        )


def _document_blocks(document: Any) -> list[tuple[str, Any]]:
    blocks: list[tuple[str, Any]] = []
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            has_image = bool(
                child.xpath(
                    ".//*[local-name()='drawing' or local-name()='pict']"
                )
            )
            blocks.append(("image_paragraph" if has_image else "paragraph", child))
        elif isinstance(child, CT_Tbl):
            blocks.append(("table", child))
    return blocks


def _is_blank_paragraph_block(kind: str, element: Any) -> bool:
    if kind != "paragraph":
        return False
    text = "".join(element.xpath(".//*[local-name()='t']/text()"))
    if text.strip():
        return False
    structural_content = element.xpath(
        ".//*[local-name()='drawing' or local-name()='pict' "
        "or local-name()='fldChar' or local-name()='instrText' "
        "or local-name()='br' or local-name()='sectPr' "
        "or local-name()='pageBreakBefore']"
    )
    return not structural_content


def _adjacent_content_kind(
    blocks: list[tuple[str, Any]],
    start_index: int,
    direction: int,
) -> str:
    index = start_index + direction
    while 0 <= index < len(blocks):
        kind, element = blocks[index]
        if not _is_blank_paragraph_block(kind, element):
            return kind
        index += direction
    return "document_boundary"


def _bookmark_names(document: Any) -> set[str]:
    return {
        name
        for name in (
            element.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name"
            )
            for element in document.element.xpath(
                ".//*[local-name()='bookmarkStart']"
            )
        )
        if name
    }


def _missing_pageref_targets(
    paragraph: Any,
    bookmark_names: set[str],
) -> list[str]:
    missing: list[str] = []
    instruction_text = "".join(
        paragraph._p.xpath(".//*[local-name()='instrText']/text()")
    )
    for match in PAGEREF_PATTERN.finditer(instruction_text):
        target = match.group(1) or match.group(2)
        if target not in bookmark_names and target not in missing:
            missing.append(target)
    return missing


def _check_caption_positions(
    factory: _IssueFactory,
    document: Any,
    analyzed: list[AnalyzedParagraph],
    rules: Rules,
) -> None:
    blocks = _document_blocks(document)
    positions = {element: index for index, (_, element) in enumerate(blocks)}
    for item in analyzed:
        if item.story != "body":
            continue
        rule = rules.roles.get(item.role)
        if rule is None or rule.expected_position is None:
            continue
        paragraph = document.paragraphs[item.paragraph_index]
        block_index = positions.get(paragraph._p)
        if block_index is None:
            continue
        if rule.expected_position == "below_object":
            direction = -1
            expected_kind = "image_paragraph"
        else:
            direction = 1
            expected_kind = "table"
        actual_kind = _adjacent_content_kind(blocks, block_index, direction)
        if actual_kind == expected_kind:
            continue
        factory.add(
            severity="warning",
            paragraph_index=item.paragraph_index,
            role=item.role,
            text_preview=item.text,
            field="expected_position",
            expected=rule.expected_position,
            actual=actual_kind,
            message=(
                f"{item.role} 未检测到要求的相邻对象位置；"
                "第一版仅报告，不自动移动。"
            ),
            fixable=False,
        )


def inspect_analyzed_document(
    document: Any,
    analyzed: list[AnalyzedParagraph],
    rules: Rules,
    document_path: str | Path,
    rules_path: str | Path,
) -> FormatReport:
    """Inspect a pre-analyzed document and return a structured report."""

    factory = _IssueFactory([])
    layout_risk = analyze_layout_risk(document, analyzed)
    _check_page_rules(factory, document, rules, layout_risk)
    ignored_toc_count = 0
    toc_unrepairable = 0
    bookmark_names = _bookmark_names(document)
    for item in analyzed:
        if item.ignore_reason == "table_of_contents":
            ignored_toc_count += 1
            if item.story != "body":
                continue
            toc_rule = resolve_role_rule(rules, item.role)
            if toc_rule is not None:
                _check_paragraph_rule(factory, item, toc_rule)
            paragraph = document.paragraphs[item.paragraph_index]
            missing_targets = _missing_pageref_targets(
                paragraph, bookmark_names
            )
            visible_field_error = bool(
                TOC_FIELD_ERROR_PATTERN.search(item.text)
            )
            if missing_targets or visible_field_error:
                repairable = (
                    bool(missing_targets)
                    and find_toc_heading_match(analyzed, item) is not None
                )
                if not repairable:
                    toc_unrepairable += 1
                factory.add(
                    severity="warning",
                    paragraph_index=item.paragraph_index,
                    role="table_of_contents",
                    text_preview=item.text,
                    field="field_result",
                    expected="有效且存在的目录书签目标",
                    actual=missing_targets or item.text,
                    message=(
                        "目录引用了不存在的 Word 书签。"
                        + (
                            "已找到唯一正文标题，可安全补建书签并请求 Word 更新域。"
                            if repairable
                            else "未找到唯一正文标题，保持不修改并列入未覆盖区域。"
                        )
                    ),
                    fixable=repairable,
                    location=item.location,
                )
            continue
        if item.role == "unknown":
            if item.ignore_reason == "protected_front_matter":
                continue
            if item.text.strip():
                factory.add(
                    severity="warning",
                    paragraph_index=item.paragraph_index,
                    role=item.role,
                    text_preview=item.text,
                    field="role",
                    expected="known role",
                    actual="unknown",
                    message="段落角色无法可靠识别，未执行强制格式检查。",
                    fixable=False,
                    location=item.location,
                )
            continue
        role_rule = resolve_role_rule(rules, item.role)
        if role_rule is not None:
            _check_paragraph_rule(factory, item, role_rule)
    _check_caption_positions(factory, document, analyzed, rules)
    reference_map = analyze_reference_document(document, analyzed)
    for problem in inspect_reference_navigation(document, reference_map):
        text_preview = (
            document.paragraphs[problem.paragraph_index].text
            if problem.paragraph_index is not None
            else ""
        )
        factory.add(
            severity=problem.severity,
            paragraph_index=problem.paragraph_index,
            role="reference",
            text_preview=text_preview,
            field=problem.field,
            expected=problem.expected,
            actual=problem.actual,
            message=problem.message,
            fixable=problem.fixable,
        )

    fixable = sum(issue.fixable for issue in factory.issues)
    summary = ReportSummary(
        total_paragraphs=len(analyzed),
        total_issues=len(factory.issues),
        fixable_issues=fixable,
        fixed_issues=0,
        unfixed_issues=len(factory.issues),
    )
    notes = [
        "段落主格式取最长非空白 run，属于可解释近似。",
        "字符缩进按代表字号换算；图题和表题位置只检测不移动。",
    ]
    if ignored_toc_count:
        notes.append(
            f"已跳过 {ignored_toc_count} 个目录段落；"
            "目录内容不参与正文格式检查或修复。"
        )
    if reference_map.entries_by_number or reference_map.citations:
        notes.append(
            f"识别到 {len(reference_map.entries_by_number)} 个参考文献编号、"
            f"{len(reference_map.citations)} 处正文引用标记。"
        )
    story_counts = Counter(item.story for item in analyzed)
    role_counts = Counter(item.role for item in analyzed)
    uncovered_areas = layout_risk.uncovered_areas()
    protected_count = sum(
        item.ignore_reason == "protected_front_matter" for item in analyzed
    )
    if protected_count:
        uncovered_areas.append(f"protected_front_matter:{protected_count}")
    if toc_unrepairable:
        uncovered_areas.append(f"toc_unrepairable:{toc_unrepairable}")
    return FormatReport(
        document=str(document_path),
        rules=str(rules_path),
        phase="inspection",
        summary=summary,
        issues=factory.issues,
        notes=notes,
        coverage=CoverageSummary(
            story_counts=dict(story_counts),
            role_counts=dict(role_counts),
            protected_front_matter=protected_count,
            layout_risk=layout_risk.to_dict(),
            uncovered_areas=uncovered_areas,
            toc_unrepairable=toc_unrepairable,
        ),
    )

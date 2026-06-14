"""Classify supported Word paragraphs into deterministic semantic roles."""

from __future__ import annotations

import re
from collections.abc import Iterator
from typing import Any

from docx.oxml.table import CT_Tbl

from .models import AnalyzedParagraph, DetectionRules, RoleName, Rules
from .ooxml_utils import get_paragraph_format

DEFAULT_PATTERNS: dict[str, list[str]] = {
    "heading_3": [r"^\d+\.\d+\.\d+(?:\s+|$)"],
    "heading_2": [r"^\d+\.\d+(?:\s+|$)"],
    "heading_1": [r"^第[一二三四五六七八九十百]+章", r"^\d+(?:\s+|$)"],
    "figure_caption": [r"^图\s*\d+", r"^Figure\s+\d+"],
    "table_caption": [r"^表\s*\d+", r"^Table\s+\d+"],
    "reference_entry": [r"^\[\d+\]", r"^\d+\.\s+"],
}

PATTERN_ORDER = (
    "title_zh",
    "title_en",
    "title",
    "abstract",
    "keywords",
    "acknowledgements",
    "appendix",
    "figure_caption",
    "table_caption",
    "heading_3",
    "heading_2",
    "heading_1",
    "reference_entry",
)

REFERENCE_HEADING_PATTERN = re.compile(
    r"^(?:参考文献|references)\s*(?:[（(].*)?$",
    flags=re.IGNORECASE,
)
REFERENCE_TERMINAL_PATTERN = re.compile(
    r"^(?:致谢|谢辞|附录|acknowledg(?:e)?ments?|appendix)(?:\s|$|[（(])",
    flags=re.IGNORECASE,
)
CJK_PATTERN = re.compile(r"[\u3400-\u4dbf\u4e00-\u9fff]")
LATIN_PATTERN = re.compile(r"[A-Za-z]")
FRONT_MATTER_METADATA_PATTERN = re.compile(
    r"(?:大学|学院|本科毕业论文|毕业设计|专业|班级|学号|学生姓名|学生|作者|"
    r"指导教师|导师|完成日期|答辩日期|thesis|university|college|author|"
    r"student|supervisor|advisor|speciality|major)",
    flags=re.IGNORECASE,
)
COVER_LABEL_PATTERN = re.compile(
    r"^(?:题目|论文题目|学院|专业|班级|学号|姓名|学生姓名|作者|指导教师|导师|"
    r"职称|日期|完成日期|英文题目|中文题目)\s*[:：]?$",
    flags=re.IGNORECASE,
)
ABSTRACT_HEADING_ZH_PATTERN = re.compile(r"^摘\s*要(?:\s*[（(].*)?$")
ABSTRACT_HEADING_EN_PATTERN = re.compile(
    r"^abstract(?:\s*[（(].*)?$",
    flags=re.IGNORECASE,
)
KEYWORDS_ZH_PATTERN = re.compile(r"^(?:关键词|关键字)\s*[:：]?", flags=re.IGNORECASE)
KEYWORDS_EN_PATTERN = re.compile(
    r"^(?:keywords?|key\s+words?)\s*(?:[（(].*?[）)])?\s*[:：]?",
    flags=re.IGNORECASE,
)
TOC_HEADING_PATTERN = re.compile(r"^(?:目\s*录|contents)$", flags=re.IGNORECASE)
TOC_ENTRY_PATTERN = re.compile(
    r"^(?P<title>.+?)(?:\.{2,}|…{2,}|\t+)\s*(?P<page>[ivxlcdm]+|\d+)\s*$",
    flags=re.IGNORECASE,
)
TOC_NUMBER_PATTERN = re.compile(r"^(?P<number>\d+(?:\.\d+){0,2})(?:\s+|$)")


def _language_role(
    text: str,
    chinese_role: RoleName,
    english_role: RoleName,
    fallback: RoleName,
) -> RoleName:
    cjk_count = len(CJK_PATTERN.findall(text))
    latin_count = len(LATIN_PATTERN.findall(text))
    if cjk_count >= 2 and cjk_count >= latin_count / 3:
        return chinese_role
    if latin_count >= 4 and cjk_count == 0:
        return english_role
    return fallback


def _title_role(text: str) -> RoleName:
    return _language_role(text, "title_zh", "title_en", "title")


def _reference_role(text: str) -> RoleName:
    return _language_role(
        text, "reference_entry_zh", "reference_entry_en", "reference_entry"
    )


def _pattern_map(detection: DetectionRules | None) -> dict[str, list[str]]:
    result = {key: list(value) for key, value in DEFAULT_PATTERNS.items()}
    if detection is None:
        return result
    explicit = detection.model_dump()
    for role in PATTERN_ORDER:
        field_name = (
            "reference_patterns"
            if role == "reference_entry"
            else f"{role}_patterns"
        )
        configured = explicit.get(field_name, [])
        if configured:
            result[role] = configured
    return result


def _role_from_patterns(
    text: str, detection: DetectionRules | None,
) -> RoleName | None:
    patterns = _pattern_map(detection)
    for role in PATTERN_ORDER:
        for pattern in patterns.get(role, []):
            if not re.search(pattern, text, flags=re.IGNORECASE):
                continue
            if role == "title":
                return _title_role(text)
            if role == "reference_entry":
                return _reference_role(text)
            return role  # type: ignore[return-value]
    return None


def _role_from_style(text: str, style_name: str) -> RoleName | None:
    normalized = style_name.strip().lower().replace("_", " ")
    compact = normalized.replace(" ", "")
    if normalized in {"title", "document title"} or compact in {"标题", "文档标题"}:
        return _title_role(text)
    for level in (1, 2, 3):
        names = {
            f"heading {level}",
            f"heading{level}",
            f"标题 {level}",
            f"标题{level}",
        }
        if normalized in names or compact in names:
            return f"heading_{level}"  # type: ignore[return-value]
    return None


def _role_from_keywords(text: str) -> RoleName | None:
    normalized = text.strip()
    lowered = normalized.lower()
    if re.match(r"^(摘要|abstract)\s*[:：]?", lowered, flags=re.IGNORECASE):
        return "abstract"
    if re.match(
        r"^(关键词|关键字|keywords?)\s*[:：]?", lowered, flags=re.IGNORECASE
    ):
        return "keywords"
    if REFERENCE_HEADING_PATTERN.match(normalized):
        return "reference_heading"
    if re.match(
        r"^(致谢|谢辞|acknowledg(?:e)?ments?)\s*$",
        normalized,
        flags=re.IGNORECASE,
    ):
        return "acknowledgements"
    if re.match(r"^(附录|appendix)(?:\s|$)", normalized, flags=re.IGNORECASE):
        return "appendix"
    if re.match(r"^(图|figure)\s*\d+", normalized, flags=re.IGNORECASE):
        return "figure_caption"
    if re.match(r"^(表|table)\s*\d+", normalized, flags=re.IGNORECASE):
        return "table_caption"
    return None


def _ignore_reason(text: str, style_name: str) -> str | None:
    normalized_style = style_name.strip().lower().replace("_", " ")
    compact_style = re.sub(r"\s+", "", normalized_style)
    is_toc_style = (
        normalized_style == "toc"
        or normalized_style.startswith("toc ")
        or bool(re.fullmatch(r"toc\d+", compact_style))
        or compact_style in {"tocheading", "目录", "目录标题"}
        or bool(re.fullmatch(r"目录\d+", compact_style))
    )
    compact_text = re.sub(r"\s+", "", text.strip()).casefold()
    if is_toc_style or compact_text in {"目录", "contents"}:
        return "table_of_contents"
    return None


def guess_paragraph_role(
    text: str,
    style_name: str,
    detection: DetectionRules | None = None,
) -> RoleName:
    """Guess one role using explicit patterns, styles, keywords, then body."""

    stripped = text.strip()
    if not stripped:
        return "unknown"
    if _ignore_reason(stripped, style_name) is not None:
        return "unknown"
    pattern_role = _role_from_patterns(stripped, detection)
    if pattern_role is not None:
        return pattern_role
    style_role = _role_from_style(stripped, style_name)
    if style_role is not None:
        return style_role
    keyword_role = _role_from_keywords(stripped)
    if keyword_role is not None:
        return keyword_role
    return "body"


def _front_matter_boundary(items: list[AnalyzedParagraph]) -> int | None:
    for item in items:
        if item.ignore_reason == "table_of_contents":
            return item.paragraph_index
        if item.role in {
            "abstract",
            "abstract_heading_zh",
            "abstract_heading_en",
            "heading_1",
            "acknowledgements",
            "appendix",
            "reference_heading",
        }:
            return item.paragraph_index
    return None


def _toc_entry_role(text: str) -> RoleName | None:
    match = TOC_ENTRY_PATTERN.match(text.strip())
    if match is None:
        return None
    number_match = TOC_NUMBER_PATTERN.match(match.group("title").strip())
    if number_match is None:
        return "toc_entry_1"
    return f"toc_entry_{min(number_match.group('number').count('.') + 1, 3)}"  # type: ignore[return-value]


def _module_role(
    text: str,
    *,
    active_abstract: str | None,
    in_toc: bool,
) -> tuple[RoleName | None, str | None, bool]:
    normalized = text.strip()
    if ABSTRACT_HEADING_ZH_PATTERN.match(normalized):
        return "abstract_heading_zh", "zh", False
    if ABSTRACT_HEADING_EN_PATTERN.match(normalized):
        return "abstract_heading_en", "en", False
    if KEYWORDS_ZH_PATTERN.match(normalized):
        return "keywords_zh", None, False
    if KEYWORDS_EN_PATTERN.match(normalized):
        return "keywords_en", None, False
    if TOC_HEADING_PATTERN.match(normalized):
        return "toc_heading", None, True
    if REFERENCE_HEADING_PATTERN.match(normalized):
        return "reference_heading", None, False
    if in_toc:
        toc_role = _toc_entry_role(normalized)
        if toc_role is not None:
            return toc_role, None, True
        if (
            "目录" in normalized
            or "contents" in normalized.casefold()
            or normalized in {"……", "..."}
        ):
            return "table_of_contents", None, True
        in_toc = False
    if active_abstract == "zh" and normalized:
        return "abstract_body_zh", active_abstract, in_toc
    if active_abstract == "en" and normalized:
        return "abstract_body_en", active_abstract, in_toc
    return None, active_abstract, in_toc


def _apply_contextual_title_detection(
    items: list[AnalyzedParagraph], protect_front_matter: bool,
) -> None:
    boundary = _front_matter_boundary(items)
    if boundary is None or boundary <= 0:
        return
    candidates = [
        item
        for item in items[:boundary]
        if item.role == "body"
        and item.text.strip()
        and not FRONT_MATTER_METADATA_PATTERN.search(item.text)
    ]
    chinese = [
        item
        for item in candidates
        if len(CJK_PATTERN.findall(item.text)) >= 6
        and len(item.text.strip()) <= 100
        and float(item.format.get("font_size_pt") or 0) >= 16
        and item.format.get("alignment") == "center"
    ]
    title_zh = max(
        chinese,
        key=lambda item: (
            float(item.format.get("font_size_pt") or 0),
            len(CJK_PATTERN.findall(item.text)),
        ),
        default=None,
    )
    if title_zh is not None:
        title_zh.role = "title_zh"

    english = [
        item
        for item in candidates
        if (title_zh is None or item.paragraph_index > title_zh.paragraph_index)
        and len(LATIN_PATTERN.findall(item.text)) >= 12
        and not CJK_PATTERN.search(item.text)
        and 10 <= len(item.text.strip()) <= 240
        and float(item.format.get("font_size_pt") or 0) >= 14
    ]
    title_en = min(
        english,
        key=lambda item: (
            item.paragraph_index - (
                title_zh.paragraph_index if title_zh is not None else 0
            ),
            -float(item.format.get("font_size_pt") or 0),
        ),
        default=None,
    )
    if title_en is not None:
        title_en.role = "title_en"

    if protect_front_matter and (title_zh is not None or title_en is not None):
        for item in items[:boundary]:
            if item.role == "body":
                item.role = "unknown"
                item.ignore_reason = "protected_front_matter"


def _iter_table_paragraphs(
    document: Any,
) -> Iterator[tuple[int, int, int, int, Any]]:
    seen_cells: set[Any] = set()
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                if cell._tc in seen_cells:
                    continue
                seen_cells.add(cell._tc)
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    yield table_index, row_index, cell_index, paragraph_index, paragraph


def _append_scoped_paragraphs(
    document: Any, analyzed: list[AnalyzedParagraph],
) -> None:
    boundary = _front_matter_boundary(
        [item for item in analyzed if item.story == "body"]
    )
    if boundary is not None:
        boundary_element = document.paragraphs[boundary]._p
    else:
        first_body = next(
            (
                item
                for item in analyzed
                if item.story == "body" and item.text.strip()
            ),
            None,
        )
        boundary_element = (
            document.paragraphs[first_body.paragraph_index]._p
            if first_body is not None
            else None
        )
    front_matter_tables: set[Any] = set()
    for child in document.element.body.iterchildren():
        if child is boundary_element:
            break
        if isinstance(child, CT_Tbl):
            front_matter_tables.add(child)

    for table_index, row_index, cell_index, paragraph_index, paragraph in (
        _iter_table_paragraphs(document)
    ):
        if not paragraph.text.strip():
            continue
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if document.tables[table_index]._tbl in front_matter_tables:
            role: RoleName = (
                "cover_label"
                if COVER_LABEL_PATTERN.match(paragraph.text.strip())
                else "cover_value"
            )
        else:
            role = "table_text"
        analyzed.append(
            AnalyzedParagraph(
                paragraph_index=paragraph_index,
                text=paragraph.text,
                role=role,
                style_name=style_name,
                format=get_paragraph_format(paragraph),
                story="table_cell",
                location=(
                    f"table[{table_index}].row[{row_index}].cell[{cell_index}]"
                    f".paragraph[{paragraph_index}]"
                ),
            )
        )

    seen_parts: set[Any] = set()
    for section_index, section in enumerate(document.sections):
        for story, part in (("header", section.header), ("footer", section.footer)):
            if part._element in seen_parts:
                continue
            seen_parts.add(part._element)
            for paragraph_index, paragraph in enumerate(part.paragraphs):
                if not paragraph.text.strip():
                    continue
                style_name = paragraph.style.name if paragraph.style is not None else ""
                analyzed.append(
                    AnalyzedParagraph(
                        paragraph_index=paragraph_index,
                        text=paragraph.text,
                        role=story,
                        style_name=style_name,
                        format=get_paragraph_format(paragraph),
                        story=story,
                        location=(
                            f"section[{section_index}].{story}"
                            f".paragraph[{paragraph_index}]"
                        ),
                    )
                )


def resolve_analyzed_paragraph(document: Any, item: AnalyzedParagraph) -> Any:
    """Resolve an analyzed paragraph back to its python-docx object."""

    if item.story == "body":
        return document.paragraphs[item.paragraph_index]
    for table_index, row_index, cell_index, paragraph_index, paragraph in (
        _iter_table_paragraphs(document)
    ):
        location = (
            f"table[{table_index}].row[{row_index}].cell[{cell_index}]"
            f".paragraph[{paragraph_index}]"
        )
        if location == item.location:
            return paragraph
    for section_index, section in enumerate(document.sections):
        part = section.header if item.story == "header" else section.footer
        for paragraph_index, paragraph in enumerate(part.paragraphs):
            location = (
                f"section[{section_index}].{item.story}.paragraph[{paragraph_index}]"
            )
            if location == item.location:
                return paragraph
    raise IndexError(f"无法解析段落位置: {item.location}")


def analyze_document(document: Any, rules: Rules) -> list[AnalyzedParagraph]:
    """Analyze body, table-cell, header, and footer paragraphs."""

    analyzed: list[AnalyzedParagraph] = []
    in_reference_section = False
    active_abstract: str | None = None
    in_toc = False
    for index, paragraph in enumerate(document.paragraphs):
        style_name = paragraph.style.name if paragraph.style is not None else ""
        ignore_reason = _ignore_reason(paragraph.text, style_name)
        text = paragraph.text.strip()
        module_role, next_abstract, next_toc = _module_role(
            text,
            active_abstract=active_abstract,
            in_toc=in_toc,
        )
        if module_role is not None:
            role = module_role
            active_abstract = next_abstract
            in_toc = next_toc
            if role in {"toc_heading", "toc_entry_1", "toc_entry_2", "toc_entry_3"}:
                ignore_reason = "table_of_contents"
        else:
            active_abstract = next_abstract
            in_toc = next_toc
            role = guess_paragraph_role(paragraph.text, style_name, rules.detection)
            if (
                ignore_reason == "table_of_contents"
                and "table_of_contents" in rules.roles
            ):
                role = "table_of_contents"
        if role == "reference_heading":
            in_reference_section = True
            active_abstract = None
            in_toc = False
        elif in_reference_section and REFERENCE_TERMINAL_PATTERN.match(text):
            in_reference_section = False
            role = _role_from_keywords(text) or role
        elif in_reference_section and text and ignore_reason is None:
            role = _reference_role(text)
        analyzed.append(
            AnalyzedParagraph(
                paragraph_index=index,
                text=paragraph.text,
                role=role,
                style_name=style_name,
                format=get_paragraph_format(paragraph),
                ignore_reason=ignore_reason,
                story="body",
                location=f"body.paragraph[{index}]",
            )
        )
    _apply_contextual_title_detection(
        analyzed, protect_front_matter=rules.priority.protect_front_matter
    )
    _append_scoped_paragraphs(document, analyzed)
    return analyzed

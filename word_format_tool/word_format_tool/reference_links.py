"""Analyze and add deterministic reference navigation to Word documents."""

from __future__ import annotations

import re
from collections import Counter, defaultdict
from copy import deepcopy
from dataclasses import dataclass
from typing import Any

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .models import AnalyzedParagraph

REFERENCE_ENTRY_PATTERN = re.compile(r"^\s*[［\[](?P<number>\d+)[］\]]")
CITATION_PATTERN = re.compile(
    r"(?:\[(?P<ascii>\d+(?:\s*[,，\-–—]\s*\d+)*)\]"
    r"|［(?P<wide>\d+(?:\s*[,，\-–—]\s*\d+)*)］)"
)
RANGE_SEPARATOR_PATTERN = re.compile(r"\s*[\-–—]\s*")
LIST_SEPARATOR_PATTERN = re.compile(r"\s*[,，]\s*")


@dataclass(frozen=True)
class CitationNumberToken:
    number: int
    start: int
    end: int


@dataclass(frozen=True)
class CitationMarker:
    text: str
    start: int
    end: int
    numbers: tuple[int, ...]
    visible_numbers: tuple[int, ...]
    number_tokens: tuple[CitationNumberToken, ...]


@dataclass(frozen=True)
class ReferenceEntry:
    number: int
    paragraph_index: int


@dataclass(frozen=True)
class CitationOccurrence:
    paragraph_index: int
    marker: CitationMarker


@dataclass(frozen=True)
class ReferenceDocument:
    entries_by_number: dict[int, ReferenceEntry]
    duplicate_entry_numbers: tuple[int, ...]
    missing_entry_numbers: tuple[int, ...]
    citations: tuple[CitationOccurrence, ...]
    cited_numbers: tuple[int, ...]
    unresolved_citation_numbers: tuple[int, ...]
    uncited_entry_numbers: tuple[int, ...]


@dataclass(frozen=True)
class ReferenceLinkResult:
    changed: bool
    linked_citations: int
    reference_bookmarks: int
    backlinks: int


@dataclass(frozen=True)
class ReferenceProblem:
    severity: str
    paragraph_index: int | None
    field: str
    expected: Any
    actual: Any
    message: str
    fixable: bool


@dataclass(frozen=True)
class _TextSlice:
    start: int
    end: int
    run: Any
    anchor: str | None


@dataclass(frozen=True)
class _DirectRunSlice:
    start: int
    end: int
    run: Any
    text: str


def reference_bookmark_name(number: int) -> str:
    return f"_SDAU_REF_{number:04d}"


def citation_bookmark_name(number: int, occurrence: int) -> str:
    return f"_SDAU_CITE_{number:04d}_{occurrence:03d}"


def _expand_citation_body(body: str) -> tuple[int, ...]:
    numbers: list[int] = []
    for part in LIST_SEPARATOR_PATTERN.split(body):
        if not part:
            continue
        endpoints = RANGE_SEPARATOR_PATTERN.split(part)
        if len(endpoints) == 1:
            numbers.append(int(endpoints[0]))
            continue
        start, end = (int(endpoints[0]), int(endpoints[-1]))
        step = 1 if end >= start else -1
        numbers.extend(range(start, end + step, step))
    return tuple(dict.fromkeys(numbers))


def find_citation_markers(text: str) -> list[CitationMarker]:
    """Return supported numeric citation markers in visible text order."""

    markers: list[CitationMarker] = []
    for match in CITATION_PATTERN.finditer(text):
        body = match.group("ascii") or match.group("wide")
        tokens = tuple(
            CitationNumberToken(
                number=int(token.group()),
                start=token.start(),
                end=token.end(),
            )
            for token in re.finditer(r"\d+", match.group())
        )
        markers.append(
            CitationMarker(
                text=match.group(),
                start=match.start(),
                end=match.end(),
                numbers=_expand_citation_body(body),
                visible_numbers=tuple(token.number for token in tokens),
                number_tokens=tokens,
            )
        )
    return markers


def _paragraph_text(paragraph: Any) -> str:
    return "".join(paragraph._p.xpath(".//*[local-name()='t']/text()"))


def analyze_reference_document(
    document: Any,
    analyzed: list[AnalyzedParagraph],
) -> ReferenceDocument:
    """Build a deterministic bibliography and body-citation index."""

    entries: dict[int, ReferenceEntry] = {}
    entry_counts: Counter[int] = Counter()
    citations: list[CitationOccurrence] = []
    roles = {
        item.paragraph_index: item.role
        for item in analyzed
        if item.story == "body"
    }

    for item in analyzed:
        if item.story != "body" or item.role not in {
            "reference_entry",
            "reference_entry_zh",
            "reference_entry_en",
        }:
            continue
        match = REFERENCE_ENTRY_PATTERN.match(item.text)
        if match is None:
            continue
        number = int(match.group("number"))
        entry_counts[number] += 1
        entries.setdefault(
            number,
            ReferenceEntry(number=number, paragraph_index=item.paragraph_index),
        )

    for index, paragraph in enumerate(document.paragraphs):
        if roles.get(index) in {
            "reference_heading",
            "reference_entry",
            "reference_entry_zh",
            "reference_entry_en",
            "unknown",
        }:
            continue
        for marker in find_citation_markers(_paragraph_text(paragraph)):
            citations.append(
                CitationOccurrence(paragraph_index=index, marker=marker)
            )

    entry_numbers = set(entries)
    cited_numbers = {
        number
        for occurrence in citations
        for number in occurrence.marker.numbers
    }
    if entry_numbers:
        missing_entries = tuple(
            number
            for number in range(1, max(entry_numbers) + 1)
            if number not in entry_numbers
        )
    else:
        missing_entries = ()

    return ReferenceDocument(
        entries_by_number=entries,
        duplicate_entry_numbers=tuple(
            sorted(number for number, count in entry_counts.items() if count > 1)
        ),
        missing_entry_numbers=missing_entries,
        citations=tuple(citations),
        cited_numbers=tuple(sorted(cited_numbers)),
        unresolved_citation_numbers=tuple(sorted(cited_numbers - entry_numbers)),
        uncited_entry_numbers=tuple(sorted(entry_numbers - cited_numbers)),
    )


def _bookmark_names(document: Any) -> set[str]:
    return {
        name
        for name in document.element.xpath(
            ".//*[local-name()='bookmarkStart']/@*[local-name()='name']"
        )
        if name
    }


def _hyperlink_anchors(element: Any) -> set[str]:
    return {
        anchor
        for anchor in element.xpath(
            ".//*[local-name()='hyperlink']/@*[local-name()='anchor']"
        )
        if anchor
    }


def _text_slices(paragraph: Any) -> list[_TextSlice]:
    slices: list[_TextSlice] = []
    cursor = 0
    for text_element in paragraph._p.xpath(".//*[local-name()='t']"):
        text = text_element.text or ""
        run = text_element.getparent()
        while run is not None and run.tag != qn("w:r"):
            run = run.getparent()
        hyperlink = text_element.getparent()
        while hyperlink is not None and hyperlink is not paragraph._p:
            if hyperlink.tag == qn("w:hyperlink"):
                break
            hyperlink = hyperlink.getparent()
        anchor = (
            hyperlink.get(qn("w:anchor"))
            if hyperlink is not None and hyperlink.tag == qn("w:hyperlink")
            else None
        )
        slices.append(
            _TextSlice(
                start=cursor,
                end=cursor + len(text),
                run=run,
                anchor=anchor,
            )
        )
        cursor += len(text)
    return slices


def _overlapping_slices(
    slices: list[_TextSlice],
    start: int,
    end: int,
) -> list[_TextSlice]:
    return [
        item for item in slices
        if item.end > start and item.start < end
    ]


def _run_is_superscript(run: Any) -> bool:
    if run is None:
        return False
    values = run.xpath(
        "./*[local-name()='rPr']/*[local-name()='vertAlign']"
        "/@*[local-name()='val']"
    )
    return bool(values and values[-1] == "superscript")


def inspect_reference_navigation(
    document: Any,
    reference_map: ReferenceDocument,
) -> list[ReferenceProblem]:
    """Return semantic and navigation problems without modifying the document."""

    problems: list[ReferenceProblem] = []
    bookmark_names = _bookmark_names(document)
    occurrence_counts: dict[int, int] = defaultdict(int)

    for number in reference_map.duplicate_entry_numbers:
        entry = reference_map.entries_by_number[number]
        problems.append(
            ReferenceProblem(
                severity="error",
                paragraph_index=entry.paragraph_index,
                field="duplicate_reference_number",
                expected="每个参考文献编号唯一",
                actual=number,
                message=f"参考文献编号 [{number}] 重复。",
                fixable=False,
            )
        )
    for number in reference_map.missing_entry_numbers:
        problems.append(
            ReferenceProblem(
                severity="warning",
                paragraph_index=None,
                field="missing_reference_number",
                expected="参考文献编号连续",
                actual=number,
                message=f"参考文献序列缺少编号 [{number}]。",
                fixable=False,
            )
        )
    for number in reference_map.unresolved_citation_numbers:
        occurrence = next(
            item
            for item in reference_map.citations
            if number in item.marker.numbers
        )
        problems.append(
            ReferenceProblem(
                severity="error",
                paragraph_index=occurrence.paragraph_index,
                field="unresolved_citation",
                expected=f"存在参考文献条目 [{number}]",
                actual=f"正文引用 [{number}] 无匹配条目",
                message=f"正文引用 [{number}] 没有对应的参考文献条目。",
                fixable=False,
            )
        )
    for number in reference_map.uncited_entry_numbers:
        entry = reference_map.entries_by_number[number]
        problems.append(
            ReferenceProblem(
                severity="warning",
                paragraph_index=entry.paragraph_index,
                field="uncited_reference",
                expected=f"正文引用参考文献 [{number}]",
                actual="未引用",
                message=f"参考文献 [{number}] 未在正文中引用。",
                fixable=False,
            )
        )

    for occurrence in reference_map.citations:
        paragraph = document.paragraphs[occurrence.paragraph_index]
        slices = _text_slices(paragraph)
        marker = occurrence.marker
        marker_slices = _overlapping_slices(
            slices, marker.start, marker.end
        )
        if marker_slices and not all(
            _run_is_superscript(item.run) for item in marker_slices
        ):
            problems.append(
                ReferenceProblem(
                    severity="error",
                    paragraph_index=occurrence.paragraph_index,
                    field="citation_superscript",
                    expected="superscript",
                    actual="not_superscript",
                    message=f"正文文献标注 {marker.text} 应使用上标格式。",
                    fixable=True,
                )
            )

        for number in marker.numbers:
            if number not in reference_map.entries_by_number:
                continue
            occurrence_counts[number] += 1
            citation_name = citation_bookmark_name(
                number, occurrence_counts[number]
            )
            if citation_name not in bookmark_names:
                problems.append(
                    ReferenceProblem(
                        severity="error",
                        paragraph_index=occurrence.paragraph_index,
                        field="citation_bookmark",
                        expected=citation_name,
                        actual="missing",
                        message=(
                            f"正文引用 [{number}] 缺少返回定位书签 "
                            f"{citation_name}。"
                        ),
                        fixable=True,
                    )
                )

        for token in marker.number_tokens:
            if token.number not in reference_map.entries_by_number:
                continue
            absolute_start = marker.start + token.start
            absolute_end = marker.start + token.end
            token_slices = _overlapping_slices(
                slices, absolute_start, absolute_end
            )
            expected_anchor = reference_bookmark_name(token.number)
            actual_anchors = {
                item.anchor for item in token_slices if item.anchor
            }
            if actual_anchors != {expected_anchor}:
                problems.append(
                    ReferenceProblem(
                        severity="error",
                        paragraph_index=occurrence.paragraph_index,
                        field="citation_link",
                        expected=expected_anchor,
                        actual=sorted(actual_anchors) or "missing",
                        message=(
                            f"正文引用数字 {token.number} 应跳转到 "
                            f"{expected_anchor}。"
                        ),
                        fixable=True,
                    )
                )

    cited_numbers = set(reference_map.cited_numbers)
    for number, entry in reference_map.entries_by_number.items():
        paragraph = document.paragraphs[entry.paragraph_index]
        reference_name = reference_bookmark_name(number)
        if reference_name not in bookmark_names:
            problems.append(
                ReferenceProblem(
                    severity="error",
                    paragraph_index=entry.paragraph_index,
                    field="reference_bookmark",
                    expected=reference_name,
                    actual="missing",
                    message=(
                        f"参考文献 [{number}] 缺少正文跳转目标书签 "
                        f"{reference_name}。"
                    ),
                    fixable=True,
                )
            )
        if number in cited_numbers:
            backlink = citation_bookmark_name(number, 1)
            if backlink not in _hyperlink_anchors(paragraph._p):
                problems.append(
                    ReferenceProblem(
                        severity="error",
                        paragraph_index=entry.paragraph_index,
                        field="reference_backlink",
                        expected=backlink,
                        actual="missing",
                        message=(
                            f"参考文献 [{number}] 的编号应返回正文首次引用 "
                            f"{backlink}。"
                        ),
                        fixable=True,
                    )
                )
    return problems


def _next_bookmark_id(document: Any) -> int:
    ids: list[int] = []
    for value in document.element.xpath(
        ".//*[local-name()='bookmarkStart' or local-name()='bookmarkEnd']"
        "/@*[local-name()='id']"
    ):
        try:
            ids.append(int(value))
        except (TypeError, ValueError):
            continue
    return max(ids, default=0) + 1


def _bookmark_pair(name: str, bookmark_id: int) -> tuple[Any, Any]:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    return start, end


def _plain_text_run(source_run: Any, text: str, *, superscript: bool) -> Any:
    run = OxmlElement("w:r")
    source_rpr = source_run.find(qn("w:rPr"))
    if source_rpr is not None:
        run.append(deepcopy(source_rpr))
    if superscript:
        rpr = run.get_or_add_rPr()
        for element in list(rpr.findall(qn("w:vertAlign"))):
            rpr.remove(element)
        vertical_align = OxmlElement("w:vertAlign")
        vertical_align.set(qn("w:val"), "superscript")
        rpr.append(vertical_align)
    text_element = OxmlElement("w:t")
    if text[:1].isspace() or text[-1:].isspace():
        text_element.set(
            "{http://www.w3.org/XML/1998/namespace}space",
            "preserve",
        )
    text_element.text = text
    run.append(text_element)
    return run


def _internal_hyperlink(anchor: str, run: Any) -> Any:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    hyperlink.append(run)
    return hyperlink


def _is_plain_text_run(run: Any) -> bool:
    return all(
        child.tag in {qn("w:rPr"), qn("w:t")}
        for child in run.iterchildren()
    )


def _replace_run(run: Any, replacements: list[Any]) -> None:
    parent = run.getparent()
    index = parent.index(run)
    parent.remove(run)
    for offset, replacement in enumerate(replacements):
        parent.insert(index + offset, replacement)


def _direct_run_slices(paragraph: Any) -> list[_DirectRunSlice]:
    slices: list[_DirectRunSlice] = []
    cursor = 0
    for child in paragraph._p.iterchildren():
        text = "".join(child.xpath(".//*[local-name()='t']/text()"))
        if child.tag == qn("w:r") and text:
            slices.append(
                _DirectRunSlice(
                    start=cursor,
                    end=cursor + len(text),
                    run=child,
                    text=text,
                )
            )
        cursor += len(text)
    return slices


def _set_run_superscript(run: Any) -> bool:
    if _run_is_superscript(run):
        return False
    rpr = run.get_or_add_rPr()
    for element in list(rpr.findall(qn("w:vertAlign"))):
        rpr.remove(element)
    vertical_align = OxmlElement("w:vertAlign")
    vertical_align.set(qn("w:val"), "superscript")
    rpr.append(vertical_align)
    return True


def _citation_occurrence_indices(
    reference_map: ReferenceDocument,
) -> dict[tuple[int, int, int], int]:
    counts: dict[int, int] = defaultdict(int)
    result: dict[tuple[int, int, int], int] = {}
    entries = set(reference_map.entries_by_number)
    for occurrence in reference_map.citations:
        for number in occurrence.marker.numbers:
            if number not in entries:
                continue
            counts[number] += 1
            result[
                (
                    occurrence.paragraph_index,
                    occurrence.marker.start,
                    number,
                )
            ] = counts[number]
    return result


def _marker_navigation_complete(
    paragraph: Any,
    occurrence: CitationOccurrence,
    reference_map: ReferenceDocument,
    occurrence_indices: dict[tuple[int, int, int], int],
    bookmark_names: set[str],
) -> bool:
    slices = _text_slices(paragraph)
    marker = occurrence.marker
    marker_slices = _overlapping_slices(slices, marker.start, marker.end)
    if not marker_slices or not all(
        _run_is_superscript(item.run) for item in marker_slices
    ):
        return False
    for number in marker.numbers:
        if number not in reference_map.entries_by_number:
            continue
        occurrence_index = occurrence_indices[
            (occurrence.paragraph_index, marker.start, number)
        ]
        if citation_bookmark_name(
            number, occurrence_index
        ) not in bookmark_names:
            return False
    for token in marker.number_tokens:
        if token.number not in reference_map.entries_by_number:
            continue
        token_slices = _overlapping_slices(
            slices,
            marker.start + token.start,
            marker.start + token.end,
        )
        if {
            item.anchor for item in token_slices if item.anchor
        } != {reference_bookmark_name(token.number)}:
            return False
    return True


def _marker_run_coverage(
    paragraph: Any,
    marker: CitationMarker,
) -> list[_DirectRunSlice]:
    overlapping = [
        item
        for item in _direct_run_slices(paragraph)
        if item.end > marker.start and item.start < marker.end
    ]
    if not overlapping:
        return []
    if (
        overlapping[0].start > marker.start
        or overlapping[-1].end < marker.end
    ):
        return []
    if any(
        left.end != right.start
        for left, right in zip(overlapping, overlapping[1:], strict=False)
    ):
        return []
    if not all(_is_plain_text_run(item.run) for item in overlapping):
        return []
    positions = [paragraph._p.index(item.run) for item in overlapping]
    if positions != list(range(positions[0], positions[0] + len(positions))):
        return []
    return overlapping


def _replace_citation_marker(
    paragraph: Any,
    occurrence: CitationOccurrence,
    reference_map: ReferenceDocument,
    occurrence_indices: dict[tuple[int, int, int], int],
    bookmark_names: set[str],
    next_id: list[int],
) -> int | None:
    marker = occurrence.marker
    coverage = _marker_run_coverage(paragraph, marker)
    if not coverage:
        return None

    replacements: list[Any] = []
    bookmarks_inserted = False
    linked_tokens: set[tuple[int, int]] = set()
    entries = set(reference_map.entries_by_number)
    absolute_tokens = [
        (
            token,
            marker.start + token.start,
            marker.start + token.end,
        )
        for token in marker.number_tokens
    ]

    for item in coverage:
        boundaries = {item.start, item.end}
        for boundary in (marker.start, marker.end):
            if item.start < boundary < item.end:
                boundaries.add(boundary)
        for _, token_start, token_end in absolute_tokens:
            if item.start < token_start < item.end:
                boundaries.add(token_start)
            if item.start < token_end < item.end:
                boundaries.add(token_end)
        ordered = sorted(boundaries)
        for start, end in zip(ordered, ordered[1:], strict=False):
            if not bookmarks_inserted and start >= marker.start:
                for number in marker.numbers:
                    if number not in entries:
                        continue
                    occurrence_index = occurrence_indices[
                        (
                            occurrence.paragraph_index,
                            marker.start,
                            number,
                        )
                    ]
                    name = citation_bookmark_name(number, occurrence_index)
                    if name in bookmark_names:
                        continue
                    bookmark_start, bookmark_end = _bookmark_pair(
                        name, next_id[0]
                    )
                    next_id[0] += 1
                    replacements.extend((bookmark_start, bookmark_end))
                    bookmark_names.add(name)
                bookmarks_inserted = True

            text = item.text[start - item.start:end - item.start]
            inside_marker = start >= marker.start and end <= marker.end
            run = _plain_text_run(
                item.run,
                text,
                superscript=inside_marker,
            )
            token_match = next(
                (
                    (token, token_start)
                    for token, token_start, token_end in absolute_tokens
                    if start >= token_start and end <= token_end
                ),
                None,
            )
            if token_match is None or token_match[0].number not in entries:
                replacements.append(run)
                continue
            token, token_start = token_match
            replacements.append(
                _internal_hyperlink(
                    reference_bookmark_name(token.number),
                    run,
                )
            )
            linked_tokens.add((token.number, token_start))

    insertion_index = paragraph._p.index(coverage[0].run)
    for item in coverage:
        paragraph._p.remove(item.run)
    for offset, replacement in enumerate(replacements):
        paragraph._p.insert(insertion_index + offset, replacement)
    return len(linked_tokens)


def _link_citation_runs(
    document: Any,
    reference_map: ReferenceDocument,
    next_id: list[int],
) -> tuple[bool, int]:
    occurrence_indices = _citation_occurrence_indices(reference_map)
    bookmark_names = _bookmark_names(document)
    by_paragraph: dict[int, list[CitationOccurrence]] = defaultdict(list)
    for occurrence in reference_map.citations:
        by_paragraph[occurrence.paragraph_index].append(occurrence)
    changed = False
    linked = 0

    for paragraph_index, occurrences in sorted(by_paragraph.items()):
        paragraph = document.paragraphs[paragraph_index]
        for occurrence in sorted(
            occurrences,
            key=lambda item: item.marker.start,
            reverse=True,
        ):
            if _marker_navigation_complete(
                paragraph,
                occurrence,
                reference_map,
                occurrence_indices,
                bookmark_names,
            ):
                continue
            replacement_count = _replace_citation_marker(
                paragraph,
                occurrence,
                reference_map,
                occurrence_indices,
                bookmark_names,
                next_id,
            )
            if replacement_count is not None:
                changed = True
                linked += replacement_count
                continue
            slices = _overlapping_slices(
                _text_slices(paragraph),
                occurrence.marker.start,
                occurrence.marker.end,
            )
            changed = any(
                _set_run_superscript(item.run) for item in slices
            ) or changed
    return changed, linked


def _link_reference_entries(
    document: Any,
    reference_map: ReferenceDocument,
    next_id: list[int],
    existing_bookmarks: set[str],
) -> tuple[int, int]:
    bookmarks_added = 0
    backlinks_added = 0
    cited_numbers = set(reference_map.cited_numbers)

    for number, entry in reference_map.entries_by_number.items():
        bookmark_name = reference_bookmark_name(number)
        paragraph = document.paragraphs[entry.paragraph_index]
        if bookmark_name in existing_bookmarks:
            backlink = citation_bookmark_name(number, 1)
            if (
                number in cited_numbers
                and backlink not in _hyperlink_anchors(paragraph._p)
                and _add_backlink_inside_bookmark(
                    paragraph, bookmark_name, backlink
                )
            ):
                backlinks_added += 1
            continue
        match = REFERENCE_ENTRY_PATTERN.match(_paragraph_text(paragraph))
        if match is None:
            continue
        label_start, label_end = match.span()
        paragraph_cursor = 0
        for run in list(paragraph._p.findall(qn("w:r"))):
            if not _is_plain_text_run(run):
                paragraph_cursor += len(
                    "".join(run.xpath(".//*[local-name()='t']/text()"))
                )
                continue
            text = "".join(run.xpath(".//*[local-name()='t']/text()"))
            run_start = paragraph_cursor
            run_end = run_start + len(text)
            paragraph_cursor = run_end
            if label_start < run_start or label_end > run_end:
                continue
            local_start = label_start - run_start
            local_end = label_end - run_start
            replacements: list[Any] = []
            if local_start:
                replacements.append(
                    _plain_text_run(run, text[:local_start], superscript=False)
                )
            bookmark_start, bookmark_end = _bookmark_pair(
                bookmark_name, next_id[0]
            )
            next_id[0] += 1
            replacements.append(bookmark_start)
            label_run = _plain_text_run(
                run,
                text[local_start:local_end],
                superscript=False,
            )
            if number in cited_numbers:
                replacements.append(
                    _internal_hyperlink(
                        citation_bookmark_name(number, 1),
                        label_run,
                    )
                )
                backlinks_added += 1
            else:
                replacements.append(label_run)
            replacements.append(bookmark_end)
            if local_end < len(text):
                replacements.append(
                    _plain_text_run(run, text[local_end:], superscript=False)
                )
            _replace_run(run, replacements)
            existing_bookmarks.add(bookmark_name)
            bookmarks_added += 1
            break
    return bookmarks_added, backlinks_added


def _add_backlink_inside_bookmark(
    paragraph: Any,
    bookmark_name: str,
    backlink: str,
) -> bool:
    starts = paragraph._p.xpath(
        "./*[local-name()='bookmarkStart' "
        f"and @*[local-name()='name']='{bookmark_name}']"
    )
    if not starts:
        return False
    start = starts[0]
    bookmark_id = start.get(qn("w:id"))
    children = list(paragraph._p.iterchildren())
    start_index = children.index(start)
    end_index = next(
        (
            index
            for index in range(start_index + 1, len(children))
            if children[index].tag == qn("w:bookmarkEnd")
            and children[index].get(qn("w:id")) == bookmark_id
        ),
        None,
    )
    if end_index is None:
        return False
    runs = [
        child
        for child in children[start_index + 1:end_index]
        if child.tag == qn("w:r")
    ]
    if not runs:
        return False
    positions = [children.index(run) for run in runs]
    if positions != list(range(positions[0], positions[0] + len(positions))):
        return False
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), backlink)
    hyperlink.set(qn("w:history"), "1")
    insertion_index = paragraph._p.index(runs[0])
    for run in runs:
        paragraph._p.remove(run)
        hyperlink.append(run)
    paragraph._p.insert(insertion_index, hyperlink)
    return True


def apply_reference_navigation(
    document: Any,
    reference_map: ReferenceDocument,
) -> ReferenceLinkResult:
    """Add missing citation links, superscripts, bookmarks, and backlinks."""

    existing_bookmarks = _bookmark_names(document)
    next_id = [_next_bookmark_id(document)]
    citation_changed, linked_citations = _link_citation_runs(
        document, reference_map, next_id
    )
    bookmarks_added, backlinks_added = _link_reference_entries(
        document,
        reference_map,
        next_id,
        existing_bookmarks,
    )
    changed = bool(
        citation_changed
        or linked_citations
        or bookmarks_added
        or backlinks_added
    )
    return ReferenceLinkResult(
        changed=changed,
        linked_citations=linked_citations,
        reference_bookmarks=bookmarks_added,
        backlinks=backlinks_added,
    )
